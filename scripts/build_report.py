"""ENCORE SRS 报告填充脚本 - 主程序.

阶段:
A. 封面 + 10 张模板自带表格
B. 正文段落替换 (1 引言, 2 项目综述, 3 系统体系结构, 4 功能需求, 5 非功能需求,
   6 验收标准, 7 产品提交, 8 签字)
C. 插入 24 张 UML 图片
D. 插入新表格 (系统用例一览, 详细用例, 性能, 验收等)
E. 保存
"""
from __future__ import annotations
import os
import sys
sys.path.insert(0, os.path.dirname(__file__))

from report_util import *
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ENCORE_NAME_CN = "ENCORE 票务管理系统"
ENCORE_NAME_EN = "ENCORE Ticketing Management System"
VERSION = "V1.0"
PROJECT_ID = "ENC-2024-001"
DOC_ID = "ENC-SRS-001"
CLASSIFICATION = "内部"
DATE_STR = "2026-05-13"


# ============================================================
# 阶段 A: 封面 + 表 1~10
# ============================================================

def fill_cover(doc):
    """封面 3 行替换 (项目名称 / 学校 / 日期), 其余留空."""
    paras = doc.paragraphs
    # paragraphs[3] = "项目名称" (Subtitle)
    replace_para_text(paras[3], ENCORE_NAME_CN)
    # paragraphs[4] = "软件需求规格说明" (主标题) - 保留
    # paragraphs[5] = "版本：V1.0" (Subtitle) - 保留, 已是 V1.0
    # paragraphs[8] = "团队成员：" - 保留
    # paragraphs[9] = "软件XXXX班XXXXXX软件公司"
    replace_para_text(paras[9], "信息学院 软件工程系")
    # paragraphs[10] = "二○××年×月"
    replace_para_text(paras[10], "二○二六年五月")
    # paragraphs[14..26]: 拟制/审核/标准化/会签/批准 - 用户要求留空


def fill_table_1_project(doc):
    """表 1 (3x2): 项目编号 / 文档编号 / 密级"""
    t = doc.tables[0]
    fill_cell(t.rows[0].cells[1], PROJECT_ID)
    fill_cell(t.rows[1].cells[1], DOC_ID)
    fill_cell(t.rows[2].cells[1], CLASSIFICATION)


def fill_table_2_team(doc):
    """表 2 (7x4): 团队成员 - 姓名/学号留空, 仅填角色与分工."""
    t = doc.tables[1]
    roles = [
        ("", "", "项目经理", "整体进度协调; 第 2 章项目综述"),
        ("", "", "前端开发", "第 3 章系统体系结构; 第 5 章界面与接口"),
        ("", "", "后端开发", "第 4 章功能需求 4.1~4.2; 第 5 章性能/安全"),
        ("", "", "数据库", "数据库设计与 ER 图; 第 5 章可靠性"),
        ("", "", "测试工程师", "第 6 章验收标准全部条目"),
        ("", "", "文档", "第 1/7/8 章; 全文校对"),
    ]
    for ri, vals in enumerate(roles, start=1):
        for ci, v in enumerate(vals):
            fill_cell(t.rows[ri].cells[ci], v)


def fill_table_3_modify_record(doc):
    """表 3 (22x5): 文档修改记录 - 仅写 V1.0 一行, 修改人留空."""
    t = doc.tables[2]
    fill_cell(t.rows[1].cells[0], "V1.0")
    fill_cell(t.rows[1].cells[1], "初稿创建: 完成 ENCORE 票务管理系统软件需求规格说明书首版编制")
    fill_cell(t.rows[1].cells[2], "")
    fill_cell(t.rows[1].cells[3], DATE_STR)
    fill_cell(t.rows[1].cells[4], "正式提交版")


def fill_table_4_terms(doc):
    """表 1 术语定义 - docx tables[3] (3x3)."""
    terms = [
        ("1", "可视化选座", "用户在剧院座位平面图中以图形化方式选择座位的过程, 座位分为可售 (AVAILABLE)、锁定中 (LOCKED)、已售 (SOLD)、不可售 (DISABLED) 四种状态."),
        ("2", "座位临时锁定", "用户选择座位后系统将该座位标记为 LOCKED 并启动倒计时, 在规定时间内未支付则自动释放为 AVAILABLE."),
        ("3", "订单倒计时", "待支付订单的有效期, 默认 15 分钟, 通过 Redis Key 过期触发; 倒计时归零自动取消订单并释放座位."),
        ("4", "电子票", "支付成功后系统生成的电子凭证, 包含票号、演出信息、场次时间、演出厅、座位号、票据状态与二维码模拟字段."),
        ("5", "检票核销", "检票员通过票号或二维码模拟码核验电子票合法性并允许入场的过程, 票据状态由 UNUSED 变为 CHECKED_IN."),
        ("6", "票房数据看板", "面向票务管理员的可视化运营分析页, 包含订单数、销售额、上座率、热门演出排行、座位区域热度等指标."),
        ("7", "场次座位票池", "为单个演出场次基于演出厅座位布局生成的座位副本集合 (schedule_seat), 是售票的最小可锁定单元."),
        ("8", "二维码模拟码", "电子票内的模拟二维码字符串字段, 检票端可直接读取, 不依赖真实扫码硬件."),
        ("9", "沉浸式剧院风格", "用户端 UI 采用暗红、金色、黑色为主色调的深色视觉语言, 与剧院文化氛围呼应."),
        ("10", "实时选座广播", "基于 WebSocket / STOMP 协议, 同一场次内任一用户的锁座 / 释放动作通过 /topic/schedule/{id}/seat 主题秒级推送给该场次全部在线用户, 屏幕端无需轮询即可同步显示."),
        ("11", "三维舞台视角预览", "用户在选座页面点击具体座位, 弹层基于 three.js + WebGL 渲染剧院 3D 场景与对应座位的 PerspectiveCamera 视角, 直观展示从该座位看向舞台的视野效果."),
        ("12", "拼座邀请", "用户在选座阶段生成一段可分享的邀请链接 / 二维码, 邀请人扫码或点击后进入同一场次, 联合锁定一组相邻座位, 形成一次共同订单."),
        ("13", "黑金指挥中心风格", "管理端数据看板采用全黑底 (#0F0F0F) + 暗红 (#5B0E0E) + 金色 (#C8A35E) 高亮的深色 BI 视觉语言, 辅以 Sankey 漏斗与座位热力图等高密度图表, 类似机场塔台 / 监控指挥中心."),
    ]
    t = doc.tables[3]
    while len(t.rows) < len(terms) + 1:
        t.add_row()
    for i, row_data in enumerate(terms, start=1):
        for ci, v in enumerate(row_data):
            fill_cell(t.rows[i].cells[ci], v)


def fill_table_5_abbr(doc):
    """表 2 缩写说明 - docx tables[4] (3x3)."""
    abbrs = [
        ("1", "SRS", "Software Requirements Specification (软件需求规格说明书)"),
        ("2", "UML", "Unified Modeling Language (统一建模语言)"),
        ("3", "B/S", "Browser/Server (浏览器/服务器架构)"),
        ("4", "API", "Application Programming Interface (应用程序编程接口)"),
        ("5", "CRUD", "Create / Read / Update / Delete (增删改查)"),
        ("6", "ER", "Entity-Relationship (实体关系)"),
        ("7", "SPA", "Single Page Application (单页应用)"),
        ("8", "ENC", "ENCORE 项目英文缩写代号"),
        ("9", "UC", "Use Case (用例)"),
        ("10", "TPS", "Transactions Per Second (每秒事务数)"),
        ("11", "QPS", "Queries Per Second (每秒查询数)"),
        ("12", "JDBC", "Java Database Connectivity (Java 数据库连接)"),
        ("13", "WS", "WebSocket (基于 TCP 的全双工通信协议)"),
        ("14", "STOMP", "Simple Text Oriented Messaging Protocol (面向消息的简单文本协议, WebSocket 之上的应用层协议)"),
        ("15", "DDD", "Domain-Driven Design (领域驱动设计)"),
        ("16", "SLA", "Service Level Agreement (服务等级协议)"),
        ("17", "FPS", "Frames Per Second (每秒帧数, 衡量 3D 渲染流畅度)"),
    ]
    t = doc.tables[4]
    while len(t.rows) < len(abbrs) + 1:
        t.add_row()
    for i, row_data in enumerate(abbrs, start=1):
        for ci, v in enumerate(row_data):
            fill_cell(t.rows[i].cells[ci], v)


def fill_table_6_references(doc):
    """表 3 引用文档 - docx tables[5] (3x6)."""
    refs = [
        ("1", "GB/T 8567-2006", "计算机软件文档编制规范", "2006", "2006-03-14", "国家标准化管理委员会"),
        ("2", "GB/T 11457-2006", "信息技术软件工程术语", "2006", "2006-03-14", "国家标准化管理委员会"),
        ("3", "ISO/IEC/IEEE 29148:2018", "Systems and software engineering — Requirements engineering", "2018", "2018-11-30", "ISO"),
        ("4", "ENC-FS-001", "ENCORE 票务管理系统可行性研究报告", "V1.0", "2026-04-20", "ENCORE 项目组"),
        ("5", "ENC-IL-001", "ENCORE 票务管理系统立项书", "V1.0", "2026-04-10", "ENCORE 项目组"),
    ]
    t = doc.tables[5]
    while len(t.rows) < len(refs) + 1:
        t.add_row()
    for i, row_data in enumerate(refs, start=1):
        for ci, v in enumerate(row_data):
            fill_cell(t.rows[i].cells[ci], v)


def fill_table_7_positions(doc):
    """表 4 岗位角色 - docx tables[6] (3x3)."""
    positions = [
        ("票务管理员", "剧院票务运营部", "维护剧目、演出分类、演出厅与座位布局、场次排期、订单查询、统计报表查看"),
        ("检票员", "剧院现场服务部", "在演出现场使用检票端核验电子票, 处理重复检票、作废票、非当前场次等异常"),
        ("剧院运营经理", "剧院运营部", "审阅票房数据看板, 把控售票与上座节奏 (本期由票务管理员账号兼任只读权限)"),
        ("系统管理员", "信息中心", "维护用户与角色、权限分配、操作日志查阅、系统参数配置"),
    ]
    t = doc.tables[6]
    while len(t.rows) < len(positions) + 1:
        t.add_row()
    for i, row_data in enumerate(positions, start=1):
        for ci, v in enumerate(row_data):
            fill_cell(t.rows[i].cells[ci], v)


def fill_table_8_actors(doc):
    """表 8 是模板原表 (3x3), 表头 名称/类别(主要/次要)/说明 → 用作系统参与者一览.
    模板中的 '表5 统计报表一览' 标题须改为 '表5 系统参与者一览'.
    """
    actors = [
        ("游客", "主要", "未登录用户; 可浏览演出列表、搜索、查看演出详情与场次基础信息、查看公告"),
        ("注册用户", "主要", "已注册并登录的用户; 可完成可视化选座、创建订单、模拟支付、查看我的订单与电子票"),
        ("票务管理员", "主要", "面向后台管理人员; 负责剧目、演出分类、演出厅、座位、场次、订单、公告、看板"),
        ("检票员", "主要", "演出现场工作人员; 负责通过票号或二维码模拟码核验电子票并处理异常"),
        ("系统管理员", "次要", "信息中心运维角色; 负责用户与角色管理、权限分配、操作日志、基础配置"),
    ]
    t = doc.tables[7]
    while len(t.rows) < len(actors) + 1:
        t.add_row()
    for i, row_data in enumerate(actors, start=1):
        for ci, v in enumerate(row_data):
            fill_cell(t.rows[i].cells[ci], v)


def fill_table_9_usecase_overview(doc):
    """表 6 系统用例一览 - docx tables[8] (7x4): 序号/用例名称/用例标识符/需求描述."""
    cases = [
        ("1", "用户注册", "ENC_UC_AUTH_00", "新用户填写账号、密码、手机号完成注册"),
        ("2", "用户登录", "ENC_UC_AUTH_01", "已注册用户输入账号与密码登录, 系统下发 Sa-Token 令牌"),
        ("3", "用户登出", "ENC_UC_AUTH_02", "已登录用户主动退出, 清理令牌"),
        ("4", "浏览演出列表", "ENC_UC_BROWSE_01", "游客或注册用户浏览全部正在售票的剧目"),
        ("5", "搜索演出", "ENC_UC_BROWSE_02", "按分类、关键词、时间检索剧目"),
        ("6", "查看演出详情", "ENC_UC_BROWSE_03", "查看剧目海报、简介、时长、票价区间与场次列表"),
        ("7", "查看场次", "ENC_UC_BROWSE_04", "选择具体演出时间与演出厅, 进入选座入口"),
        ("8", "可视化选座", "ENC_UC_SEAT_01", "在座位图上选择一个或多个 AVAILABLE 座位并临时锁定"),
        ("9", "创建订单", "ENC_UC_ORDER_01", "依据已锁座位生成待支付订单并启动倒计时"),
        ("10", "取消订单", "ENC_UC_ORDER_02", "用户主动取消待支付订单, 释放座位"),
        ("11", "模拟支付", "ENC_UC_PAY_01", "用户在倒计时内完成支付模拟, 订单状态 PENDING_PAYMENT→PAID"),
        ("12", "查看我的订单", "ENC_UC_ORDER_03", "用户分页查看本人的全部订单及当前状态"),
        ("13", "查看电子票", "ENC_UC_PAY_02", "支付成功后查看票号、二维码模拟码与座位信息"),
        ("14", "演出分类管理", "ENC_UC_SHOW_00", "票务管理员维护分类树, 含新增/修改/启停"),
        ("15", "剧目信息管理", "ENC_UC_SHOW_01", "票务管理员维护剧目档案与海报"),
        ("16", "演出厅管理", "ENC_UC_HALL_01", "票务管理员维护演出厅基本信息"),
        ("17", "座位布局管理", "ENC_UC_HALL_02", "票务管理员维护演出厅座位行列与区域"),
        ("18", "场次排期管理", "ENC_UC_SCH_01", "票务管理员为剧目创建场次并选定演出厅与开演时间"),
        ("19", "场次座位票池管理", "ENC_UC_SCH_02", "基于演出厅座位生成场次专属票池, 支持票价分区"),
        ("20", "订单查询与管理", "ENC_UC_ORDER_04", "票务管理员检索和处理订单, 含异常订单"),
        ("21", "公告管理", "ENC_UC_ANN_01", "票务管理员发布、修改、下线剧院公告"),
        ("22", "查看数据看板", "ENC_UC_STAT_01", "票务管理员查看销售额、订单数、上座率、热门演出等核心指标"),
        ("23", "检票核销", "ENC_UC_CHECKIN_01", "检票员通过票号或二维码模拟码核销电子票并允许入场"),
        ("24", "异常票据处理", "ENC_UC_CHECKIN_02", "检票员对重复检票、作废票、非当前场次票据进行异常处理"),
        ("25", "用户账号管理", "ENC_UC_SYS_01", "系统管理员创建、禁用、重置用户账号"),
        ("26", "角色与权限管理", "ENC_UC_SYS_02", "系统管理员维护角色清单并为用户分配角色"),
        ("27", "操作日志查看", "ENC_UC_SYS_03", "系统管理员查询关键操作日志并导出"),
    ]
    t = doc.tables[8]
    while len(t.rows) < len(cases) + 1:
        t.add_row()
    for i, row_data in enumerate(cases, start=1):
        for ci, v in enumerate(row_data):
            fill_cell(t.rows[i].cells[ci], v)


def fill_table_10_signature(doc):
    """表 10 签字: 仅替换占位单位名, 个人姓名/日期留空."""
    t = doc.tables[9]
    fill_cell(t.rows[1].cells[1], "剧院运营方")
    fill_cell(t.rows[1].cells[4], "ENCORE 项目组")


# ============================================================
# 阶段 B: 正文段落替换 (1 引言 ~ 8 签字)
# ============================================================

def _find_h(doc, text):
    """精确定位某 Heading 段落 (按文字)."""
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f"找不到 Heading: {text!r}")


def _next_anchor(doc, after, *anchors):
    """从 after 之后, 找下一个出现在 anchors 中的段落, 返回它."""
    seen = False
    for p in doc.paragraphs:
        if not seen:
            if p is after:
                seen = True
            continue
        if p.text.strip() in anchors:
            return p
    return None


def _replace_range(doc, start_heading_text: str, next_heading_text: str | None,
                   new_items, keep_tables: bool = True):
    """把 start_heading 段落后到 next_heading 段落前的所有 w:p 删除 (默认保留 w:tbl);
    然后把 new_items = [(text, style)] 在 start_heading 后插入."""
    head = _find_h(doc, start_heading_text)
    next_el = None
    if next_heading_text:
        nxt = _find_h(doc, next_heading_text)
        next_el = nxt._p
    body = doc.element.body
    # 删除 head 后到 next_el 前的所有 w:p
    el = head._p.getnext()
    while el is not None and el is not next_el:
        nxt = el.getnext()
        if el.tag == qn('w:p'):
            body.remove(el)
        elif el.tag == qn('w:tbl') and not keep_tables:
            body.remove(el)
        el = nxt
    # 在 head 后插入 new_items (倒序插入保持顺序)
    cur = head
    for text, style in new_items:
        cur = insert_para_after(cur, text, style)


def stage_b_chapter1_intro(doc):
    """1 引言 (1.1 编写目的, 1.2 项目资料, 1.3 术语定义, 1.4 缩写说明, 1.5 引用文档)."""
    # 章节首段
    _replace_range(doc, "引言", "编写目的", [
        ("本章节为《ENCORE 票务管理系统软件需求规格说明书》的引言部分, 阐述本文档的编写目的、项目基本资料、术语与缩写说明以及引用文档清单, 为读者快速建立 ENCORE 项目的整体背景认知.", "Normal"),
    ])

    # 1.1 编写目的
    _replace_range(doc, "编写目的", "项目资料", [
        ("本文档以 ENCORE 票务管理系统的需求方为视角, 对系统的功能与非功能需求进行规格化描述, 用于指导后续概要设计、详细设计、编码实现、测试验收以及上线后的运维管理.", "Normal"),
        ("本文档预期读者包括:", "Normal"),
        ("ENCORE 项目组全体成员 (项目经理、前端、后端、数据库、测试、文档);", "List Paragraph"),
        ("剧院方业务代表 (票务管理员、检票员、剧院运营经理);", "List Paragraph"),
        ("软件工程课程设计 II 指导教师与评审组;", "List Paragraph"),
        ("信息中心系统管理员及未来运维接手人员.", "List Paragraph"),
    ])

    # 1.2 项目资料
    _replace_range(doc, "项目资料", "术语定义", [
        (f"项目名称: {ENCORE_NAME_CN}", "List Paragraph"),
        (f"项目英文名: {ENCORE_NAME_EN}", "List Paragraph"),
        (f"项目编号: {PROJECT_ID}", "List Paragraph"),
        (f"文档编号: {DOC_ID}", "List Paragraph"),
        (f"版本号: {VERSION}", "List Paragraph"),
        ("架构类型: B/S 架构, 前后端分离", "List Paragraph"),
        ("项目定位: 面向中小型剧院、校园剧场和演出场馆的现代化票务运营管理系统", "List Paragraph"),
        ("覆盖业务: 演出发布 → 用户购票 → 现场检票 → 经营分析", "List Paragraph"),
        ("项目周期: 软件工程课程设计 II (2024-05-27 至 2024-06-07) 后持续迭代", "List Paragraph"),
        ("项目阶段成果: 软件需求规格说明书、概要设计说明书、详细设计说明书、源代码、测试报告、用户手册", "List Paragraph"),
    ])

    # 1.3 术语定义 - 保留表 1 (术语), 仅替换前后说明
    head = _find_h(doc, "术语定义")
    nxt = _find_h(doc, "缩写说明")
    body = doc.element.body
    body_wrap = doc._body  # 必须用 _body 作为 Paragraph 的 parent
    el = head._p.getnext()
    while el is not None and el is not nxt._p:
        nxt_el = el.getnext()
        if el.tag == qn('w:p'):
            body.remove(el)
        el = nxt_el
    head = _find_h(doc, "术语定义")
    cur = insert_para_after(head, "本文档中使用的专业术语定义见表 1.", "Normal")
    cur = insert_para_after(cur, "表1  术语定义", "077表")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl_el = doc.tables[3]._tbl
    tbl_el.getparent().remove(tbl_el)
    cur._p.addnext(tbl_el)
    new_tail = OxmlElement('w:p')
    tbl_el.addnext(new_tail)
    tail_p = Paragraph(new_tail, body_wrap)
    set_style(tail_p, "Normal")
    tail_p.add_run("如表 1 所示, 系统所涉术语紧密围绕 ENCORE 票务全生命周期管理, 既包含与用户体验相关的可视化选座、电子票、二维码模拟码, 也包含支撑业务约束的座位临时锁定、订单倒计时、检票核销等核心机制.")

    # 1.4 缩写说明 - 类似处理 (保留表 2)
    head = _find_h(doc, "缩写说明")
    nxt = _find_h(doc, "引用文档")
    el = head._p.getnext()
    while el is not None and el is not nxt._p:
        nxt_el = el.getnext()
        if el.tag == qn('w:p'):
            body.remove(el)
        el = nxt_el
    head = _find_h(doc, "缩写说明")
    cur = insert_para_after(head, "本文档中使用的英文及行业缩写见表 2.", "Normal")
    cur = insert_para_after(cur, "表2  英文缩写说明", "077表")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl_el = doc.tables[4]._tbl
    tbl_el.getparent().remove(tbl_el)
    cur._p.addnext(tbl_el)
    new_tail = OxmlElement('w:p')
    tbl_el.addnext(new_tail)
    tail_p = Paragraph(new_tail, body_wrap)
    set_style(tail_p, "Normal")
    tail_p.add_run("缩写在正文中首次出现时同时给出中文全称, 后续出现时直接使用缩写.")

    # 1.5 引用文档
    head = _find_h(doc, "引用文档")
    nxt = _find_h(doc, "项目综述")
    el = head._p.getnext()
    while el is not None and el is not nxt._p:
        nxt_el = el.getnext()
        if el.tag == qn('w:p'):
            body.remove(el)
        el = nxt_el
    head = _find_h(doc, "引用文档")
    cur = insert_para_after(head, "本文档编制过程中引用的标准、规范与项目内部资料见表 3.", "Normal")
    cur = insert_para_after(cur, "表3  引用文档", "077表")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl_el = doc.tables[5]._tbl
    tbl_el.getparent().remove(tbl_el)
    cur._p.addnext(tbl_el)
    new_tail = OxmlElement('w:p')
    tbl_el.addnext(new_tail)
    tail_p = Paragraph(new_tail, body_wrap)
    set_style(tail_p, "Normal")
    tail_p.add_run("表 3 中列出的 GB/T 8567 与 ISO/IEC/IEEE 29148 是本文档结构与术语的国际/国内标准依据; ENCORE 项目可行性研究报告与立项书是本文档的上游交付物.")

    print("[Stage B-1] 第 1 章 引言 替换完成")


def stage_b_chapter2_overview(doc):
    """2 项目综述 (2.1 项目背景, 2.2 组织机构与职责, 2.3 岗位角色, 2.4 业务流程, 2.5 统计报表)."""
    _replace_range(doc, "项目综述", "项目背景", [
        ("本章节从业务视角介绍 ENCORE 票务管理系统的项目背景、所服务的剧院组织机构与岗位角色、关键业务流程以及配套的统计报表能力, 以建立后续章节展开功能需求与非功能需求所需的业务上下文.", "Normal"),
    ])

    _replace_range(doc, "项目背景", "组织机构与职责", [
        ("近年, 中小型剧院、校园剧场以及商业演出场馆纷纷尝试将线下票务流程数字化, 但目前国内多数中小剧院仍存在以下痛点:", "Normal"),
        ("人工售票效率低: 高度依赖窗口售票, 高峰期排队严重, 错票漏票时有发生;", "List Paragraph"),
        ("座位状态不透明: 用户无法在购票前直观看到座位分布与已售情况, 体验欠佳;", "List Paragraph"),
        ("订单与检票脱节: 票面信息以纸质票为载体, 检票过程难以与售票数据实时联动;", "List Paragraph"),
        ("销售统计滞后: 票房数据无法实时反馈, 经营决策严重依赖人工汇总;", "List Paragraph"),
        ("纸质票据管理不便: 票据印制、分发、回收均产生显著耗材成本与流程负担;", "List Paragraph"),
        ("票务数据难以沉淀: 用户、订单、检票、上座率等数据散落于不同环节, 难以形成可分析的资产.", "List Paragraph"),
        ("基于以上痛点, ENCORE 票务管理系统以 \"票务生命周期管理\" 为核心理念, 通过可视化选座、座位临时锁定、订单倒计时、电子票核销和票房数据看板, 构建覆盖 \"演出发布 → 用户购票 → 现场检票 → 经营分析\" 的完整剧院票务闭环, 帮助剧院方将原本分散、滞后的票务环节整合为统一、实时、可分析的数字化运营资产.", "Normal"),
    ])

    _replace_range(doc, "组织机构与职责", "岗位角色", [
        ("ENCORE 票务管理系统服务的剧院组织通常采用如下机构划分:", "Normal"),
        ("剧院运营部: 总体管理剧目策划、场次排期、上座率与营收指标;", "List Paragraph"),
        ("剧院票务运营部: 日常维护剧目档案、演出厅与座位、场次售票状态、订单查询;", "List Paragraph"),
        ("剧院现场服务部: 演出现场负责检票核销、座席引导、特殊票据异常处置;", "List Paragraph"),
        ("信息中心: 维护账号体系、权限分配、操作日志、系统运维.", "List Paragraph"),
        ("系统的权限模型与上述机构对齐: 票务管理员对应运营/票务运营部, 检票员对应现场服务部, 系统管理员对应信息中心. 剧院运营经理在本期作为只读角色, 通过票务管理员账号查看数据看板.", "Normal"),
    ])

    _replace_range(doc, "岗位角色", "业务流程", [
        ("ENCORE 票务管理系统涉及的主要岗位角色见表 4. 后续功能需求中的参与者定义与此保持一致.", "Normal"),
        ("表4  岗位角色", "077表"),
    ])
    # 在表 4 标题段后插入表 7 (岗位角色)
    cap = _find_h(doc, "表4  岗位角色")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl_el = doc.tables[6]._tbl
    tbl_el.getparent().remove(tbl_el)
    cap._p.addnext(tbl_el)
    new_tail = OxmlElement('w:p')
    tbl_el.addnext(new_tail)
    tp = Paragraph(new_tail, doc._body)
    set_style(tp, "Normal")
    tp.add_run("岗位角色侧重组织/部门视角, 与第 4.1 节系统参与者一一映射但描述维度不同: 岗位角色描述部门职责, 参与者定义描述其与系统的交互能力.")

    _replace_range(doc, "业务流程", "统计报表", [
        ("ENCORE 票务管理系统的核心业务流程围绕 \"剧目—场次—座位票池—订单—电子票—检票—统计\" 建立完整闭环, 总流程如图 9 所示, 其文字描述如下:", "Normal"),
        ("图9  ENCORE 核心业务总活动图", "077图"),
        ("① 票务管理员在后台创建演出分类与剧目档案, 上传剧目海报、简介、时长等基本信息;", "List Paragraph"),
        ("② 票务管理员维护演出厅信息以及演出厅的座位行列与区域布局 (hall_seat 表);", "List Paragraph"),
        ("③ 票务管理员为剧目创建具体演出场次, 指定演出厅与开演时间, 并设置场次票价分区;", "List Paragraph"),
        ("④ 系统依据演出厅座位布局自动生成该场次的座位票池 (schedule_seat 表, 全部 AVAILABLE);", "List Paragraph"),
        ("⑤ 场次审核通过后, 票务管理员将场次状态从 DRAFT 切换为 ON_SALE 开始售票;", "List Paragraph"),
        ("⑥ 注册用户登录后浏览演出, 进入演出详情页, 选定具体场次进入可视化选座页;", "List Paragraph"),
        ("⑦ 用户在座位图中选择一个或多个 AVAILABLE 座位, 系统通过 Redis 临时锁定座位 (LOCKED) 并创建待支付订单 (PENDING_PAYMENT), 启动默认 15 分钟订单倒计时;", "List Paragraph"),
        ("⑧ 用户在倒计时窗口内提交模拟支付, 订单状态变为 PAID, 座位状态变为 SOLD, 系统批量生成电子票 (ticket_item, UNUSED) 并下发票号与二维码模拟码;", "List Paragraph"),
        ("⑨ 若用户未在倒计时内完成支付, 订单自动变为 EXPIRED, 系统释放对应座位 (LOCKED → AVAILABLE);", "List Paragraph"),
        ("⑩ 演出当日, 检票员在检票端输入或扫描票号, 系统核验票据状态、场次时间, 通过则票据置为 CHECKED_IN 并写入 checkin_record;", "List Paragraph"),
        ("⑪ 演出结束后, 票务管理员通过数据看板查看销售额、订单数、上座率、热门演出排行等指标, 支撑后续排期与运营决策.", "List Paragraph"),
        ("该业务闭环既是后续 4.2 节用例图与用例一览表的设计依据, 也是 4.3 节核心用例详述、状态机以及 5.2 节性能指标的核心场景.", "Normal"),
    ])
    cap9 = _find_h(doc, "图9  ENCORE 核心业务总活动图")
    cap9.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _replace_range(doc, "统计报表", "系统体系结构", [
        ("ENCORE 票务管理系统的统计报表能力面向票务管理员的票房数据看板, 用于辅助剧院经营分析与排期决策, 报表清单见表 5.", "Normal"),
        ("表5  系统统计报表一览", "077表"),
    ], keep_tables=False)
    cap = _find_h(doc, "表5  系统统计报表一览")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 新建一个 3x3 报表表 (因为模板的 Table 8 我们已挪用为参与者表, 而且参与者要在 4.1)
    # 实际上: tables[7] 还在原位, 即 4.1 节. 我们这里要新建.
    tbl = insert_table_after(cap, rows=9, cols=3)
    fill_row(tbl.rows[0], ["报表名称", "类别 (主要/次要)", "说明"], bold=True, center=True)
    rows_data = [
        ("今日订单数", "主要", "按支付日期统计当日成功支付的订单总数"),
        ("今日销售额", "主要", "按支付日期统计当日成功支付订单的成交金额合计"),
        ("近 7 日票房趋势", "主要", "近 7 个自然日每日订单数与销售额的折线趋势"),
        ("热门演出 TOP 10", "主要", "按近 30 日销售额对剧目排行"),
        ("场次上座率", "主要", "每场次 已售座位 / 总座位 的比率"),
        ("演出厅利用率", "次要", "每演出厅 已开场次 / 可排场次 的比率"),
        ("订单状态分布", "次要", "PENDING_PAYMENT / PAID / EXPIRED / CANCELLED / REFUNDED 占比饼图"),
        ("座位区域热度", "次要", "按座位区域统计已售比率, 用于优化票价分区"),
    ]
    for i, r in enumerate(rows_data, start=1):
        fill_row(tbl.rows[i], list(r))
    new_tail = OxmlElement('w:p')
    tbl._tbl.addnext(new_tail)
    tp = Paragraph(new_tail, doc._body)
    set_style(tp, "Normal")
    tp.add_run("表 5 中 \"主要\" 类报表是数据看板首屏的核心指标; \"次要\" 类报表在二级页面展开, 用于细粒度的经营分析.")
    print("[Stage B-2] 第 2 章 项目综述 替换完成")


def stage_b_chapter3_arch(doc):
    """3 系统体系结构 (3.1 物理架构, 3.2 网络环节, 3.3 运行环境)."""
    _replace_range(doc, "系统体系结构", "物理架构", [
        ("本章节描述 ENCORE 票务管理系统的物理架构、网络环节与运行环境. 系统采用 B/S 架构与前后端分离设计, 业务逻辑全部集中于后端, 通过浏览器交付前端 SPA, 部署成本与运维复杂度低, 适合中小型剧院场景.", "Normal"),
    ])

    _replace_range(doc, "物理架构", "网络环节", [
        ("ENCORE 票务管理系统的物理架构如图 1 所示. 用户、票务管理员、检票员均通过浏览器访问同一套前端 SPA, 由 Nginx 负责静态资源与 /api/** 反向代理, Spring Boot 应用承载全部业务逻辑, MySQL 与 Redis 分别承担持久化与短期状态.", "Normal"),
        ("图1  ENCORE 票务管理系统物理架构", "077图"),
        ("如图 1 所示, ENCORE 在物理上由用户终端浏览器、应用服务器 (Nginx + Spring Boot)、数据服务 (MySQL + Redis) 三层构成. MySQL 持久化全部业务数据; Redis 承担登录态、座位临时锁、订单倒计时等高频读写但短生命周期的数据, 显著降低对 MySQL 的并发压力.", "Normal"),
    ])

    _replace_range(doc, "网络环节", "运行环境", [
        ("ENCORE 票务管理系统的网络拓扑如图 2 所示. 外网入口仅暴露 Nginx 的 80/443 端口; 应用服务器、数据库服务器、缓存服务器之间通过内网通信, 减小攻击面, 也降低跨网延迟. 检票员可通过演出现场的 Wi-Fi 局域网或商用 4G/5G 接入.", "Normal"),
        ("图2  ENCORE 票务管理系统网络拓扑", "077图"),
        ("如图 2 所示, 演出现场的检票端、剧院后台的票务管理端与互联网中的用户购票端共享同一套后端 API, 但通过 Sa-Token 角色与权限差异加以隔离, 不允许检票员或管理员账号在公网无 IP 白名单情况下访问敏感接口.", "Normal"),
    ])

    _replace_range(doc, "运行环境", "功能需求", [
        ("ENCORE 票务管理系统的运行环境配置如下:", "Normal"),
        ("服务器操作系统: 兼容 Linux (Ubuntu 22.04 LTS / CentOS Stream 9) 与 Windows Server 2019 及以上;", "List Paragraph"),
        ("JDK: OpenJDK 17 或 21, 长期支持版本;", "List Paragraph"),
        ("Web 服务器: Nginx 1.24+, 负责静态资源与反向代理 (含 WebSocket Upgrade 转发);", "List Paragraph"),
        ("数据库: MySQL 8.0+ (utf8mb4 字符集), 推荐部署在独立服务器或托管实例;", "List Paragraph"),
        ("缓存: Redis 7.0+, 启用持久化, 用于座位锁与订单倒计时, 并兼作实时选座广播的 STOMP relay;", "List Paragraph"),
        ("浏览器兼容性: 用户端与管理端兼容 Chrome 100+ / Edge 100+ / Firefox 100+ / Safari 15+, 不再支持 IE;", "List Paragraph"),
        ("终端分辨率: 用户端不低于 1280×720, 管理端不低于 1366×768, 检票端不低于 800×600 (横屏);", "List Paragraph"),
        ("前端技术栈: Vue 3 + TypeScript + Vite + Vue Router + Pinia + Element Plus + ECharts + Axios + SCSS; 引入 three.js 0.16x 渲染剧院 3D 视角, @stomp/stompjs 作为 WebSocket 客户端;", "List Paragraph"),
        ("后端技术栈: Spring Boot 3 + Java 17 + MyBatis-Plus + Sa-Token + Spring WebSocket (simp messaging) + Knife4j 4.x (Swagger UI) + Lombok + Maven;", "List Paragraph"),
        ("部署工具: Docker Compose 编排; GitHub 作为代码托管, DataGrip 作为数据库 IDE, Apifox/Postman 用于接口联调; JMeter 5.6 + Skywalking 9.x 用于全链路压测与调用链追踪.", "List Paragraph"),
    ])
    print("[Stage B-3] 第 3 章 系统体系结构 替换完成")


# ============================================================
# 阶段 B-4: 第 4 章 功能需求 (大块)
# ============================================================

def _detailed_uc_rows(uc_id, name, scope, level, primary_actor, stakeholders,
                     pre, post, main_flow, ext_flow, special, tech_var):
    """生成 12 行详细用例说明表的内容."""
    return [
        ["用例标识符", uc_id],
        ["用例名称", name],
        ["范围", scope],
        ["级别", level],
        ["主要角色", primary_actor],
        ["涉众及关注点", stakeholders],
        ["前置条件", pre],
        ["后置条件", post],
        ["主成功场景", main_flow],
        ["扩展或替代流程", ext_flow],
        ["特殊需求", special],
        ["技术和数据变元表", tech_var],
    ]


def _insert_detail_uc(after_para, title_text, table_caption, rows, doc):
    """在 after_para 后插入: 子节标题段 + 表标题段 + 12 行详细表 + 后段."""
    head = insert_para_after(after_para, title_text, "Heading 3")
    cap = insert_para_after(head, table_caption, "077表")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl = insert_table_after(cap, rows=len(rows), cols=2)
    # 设置首列加粗
    for i, (k, v) in enumerate(rows):
        fill_cell(tbl.rows[i].cells[0], k, bold=True)
        fill_cell(tbl.rows[i].cells[1], v)
    # 后缀段
    tail_el = OxmlElement('w:p')
    tbl._tbl.addnext(tail_el)
    tail_p = Paragraph(tail_el, doc._body)
    set_style(tail_p, "Normal")
    return tail_p


def stage_b_chapter4_functional(doc):
    """4 功能需求 - 完整重建.

    最终结构:
      4 功能需求 (引言段)
      4.1 参与者定义 - 文本 + 表6 系统参与者一览 + 图3 总用例图 + 图4~8 子用例图
      4.2 功能构成 - 文本 + 图20 类图 + 图24 ER 图 + 表7 系统用例一览 + 表8 数据库实体说明
      4.3 ENC_UC_AUTH_01   用户注册与登录 (表9 + 图14)
      4.4 ENC_UC_BROWSE_01 浏览与查询演出 (表10)
      4.5 ENC_UC_SEAT_01   可视化选座 (表11 + 图10/15/21)
      4.6 ENC_UC_ORDER_01  创建订单 (表12 + 图16/22)
      4.7 ENC_UC_PAY_01    模拟支付与电子票生成 (表13 + 图11/17/23)
      4.8 ENC_UC_CHECKIN_01检票核销 (表14 + 图12/18)
      4.9 ENC_UC_SHOW_01   剧目与场次管理 (表15 + 图13)
      4.10 ENC_UC_STAT_01  统计报表查看 (表16 + 图19)
    """
    # ---- 全章删除 (含模板原 4.1 参与者表) + 重建 ----
    head = _find_h(doc, "功能需求")
    next_h = _find_h(doc, "非功能需求")
    body = doc.element.body
    body_wrap = doc._body
    el = head._p.getnext()
    while el is not None and el is not next_h._p:
        nxt = el.getnext()
        if el.tag in (qn('w:p'), qn('w:tbl')):
            body.remove(el)
        el = nxt

    # 章节首段
    cur = insert_para_after(head, "本章节描述 ENCORE 票务管理系统的功能需求, 包括系统参与者、功能构成、核心用例的详细说明, 以及配套的核心领域类图、数据库实体关系图.", "Normal")

    # ============ 4.1 参与者定义 ============
    cur = insert_para_after(cur, "参与者定义", "Heading 2")
    cur = insert_para_after(cur, "参与者 (Actor) 指与 ENCORE 票务管理系统产生交互的外部用户或外部系统. 依据 ENCORE 项目差异化定位, 系统共定义 5 类参与者, 其角色与系统的交互能力见表 6.", "Normal")
    cur = insert_para_after(cur, "表6  ENCORE 系统参与者一览", "077表")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    actors = [
        ("ENC_ACTOR_01", "游客", "主要", "未登录用户, 可浏览演出列表、搜索、查看演出详情与场次基础信息、查看公告. 不能选座、下单和查看电子票."),
        ("ENC_ACTOR_02", "注册用户", "主要", "已注册并登录的最终消费者. 可执行可视化选座、创建订单、模拟支付、查看我的订单、查看电子票与二维码模拟码等核心购票动作."),
        ("ENC_ACTOR_03", "票务管理员", "主要", "剧院票务运营部岗位. 负责剧目、演出分类、演出厅、座位布局、场次、座位票池、订单、公告、数据看板的全部后台管理."),
        ("ENC_ACTOR_04", "检票员", "主要", "剧院现场服务部岗位. 负责通过票号或二维码模拟码核销电子票, 并处理重复检票、作废票、非当前场次等异常."),
        ("ENC_ACTOR_05", "系统管理员", "次要", "信息中心岗位. 负责用户与角色管理、权限分配、操作日志、基础参数配置. 不参与剧目和票务业务."),
    ]
    actor_tbl = insert_table_after(cur, rows=len(actors) + 1, cols=5)
    fill_row(actor_tbl.rows[0], ["编号", "参与者名称", "类别", "标识符", "描述"], bold=True, center=True)
    for i, (sid, name, kind, desc) in enumerate(actors, start=1):
        fill_row(actor_tbl.rows[i], [str(i), name, kind, sid, desc])
    cur_el = OxmlElement('w:p')
    actor_tbl._tbl.addnext(cur_el)
    cur = Paragraph(cur_el, body_wrap)
    set_style(cur, "Normal")
    cur.add_run("注: 主要参与者直接从系统获取业务价值; 次要参与者为系统提供支撑性服务.")

    cur = insert_para_after(cur, "ENCORE 票务管理系统的总用例图如图 3 所示, 5 类参与者与系统的关键用例一览无遗.", "Normal")
    cur = insert_para_after(cur, "图3  ENCORE 票务管理系统总用例图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cur = insert_para_after(cur, "为便于聚焦各类参与者, 下面分别给出游客、注册用户、票务管理员、检票员、系统管理员对应的子用例图.", "Normal")
    cur = insert_para_after(cur, "图4  游客浏览演出用例图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur = insert_para_after(cur, "图5  注册用户购票用例图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur = insert_para_after(cur, "图6  票务管理员后台管理用例图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur = insert_para_after(cur, "图7  检票员检票核销用例图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur = insert_para_after(cur, "图8  系统管理员权限管理用例图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ============ 4.2 功能构成 ============
    cur = insert_para_after(cur, "功能构成", "Heading 2")
    cur = insert_para_after(cur, "ENCORE 票务管理系统从用户端、管理端、检票端三个使用界面切入, 整体功能模块可归纳为 12 大类, 共 27 个核心用例 (见表 7). 系统的功能模块构成如下:", "Normal")
    for f in [
        "用户端: 用户注册与登录、演出浏览与搜索、演出详情查看、场次选择、可视化选座、订单创建与模拟支付、我的订单、我的电子票;",
        "管理端: 用户与角色管理、剧目信息管理、演出分类管理、演出厅管理、座位布局管理、场次排期管理、场次座位票池管理、订单管理、检票管理 (二级菜单概览)、数据统计看板、公告管理、操作日志管理;",
        "检票端: 票号 / 二维码模拟码输入、票据核验、检票入场、异常处理.",
    ]:
        cur = insert_para_after(cur, f, "List Paragraph")

    cur = insert_para_after(cur, "为方便后续设计阶段的类与表对齐, 本节先给出 ENCORE 核心领域类图 (图 20) 与数据库实体关系图 (图 24), 再列出系统全部用例.", "Normal")
    cur = insert_para_after(cur, "图20  ENCORE 核心领域类图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur = insert_para_after(cur, "如图 20 所示, ENCORE 的核心领域对象包括 User / Role / UserRole 三类账号实体, ShowCategory / ShowInfo 两类剧目实体, TheaterHall / HallSeat 两类演出厅实体, ShowSchedule / ScheduleSeat 两类场次实体, TicketOrder / TicketItem 两类订单与票据实体, 以及 PaymentRecord / CheckinRecord 两类辅助记录实体. 状态枚举 ScheduleStatus / SeatStatus / OrderStatus / TicketStatus 与正文及状态图保持一致.", "Normal")
    cur = insert_para_after(cur, "图24  ENCORE 数据库实体关系图 (ER)", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur = insert_para_after(cur, "如图 24 所示, ENCORE 在数据库层面共定义 15 张核心表. 数据库实体的字段语义与关系说明见表 8.", "Normal")

    # 表 7 系统用例一览 - 内联创建
    cur = insert_para_after(cur, "系统全部用例的标识符、名称与功能描述见表 7.", "Normal")
    cur = insert_para_after(cur, "表7  ENCORE 系统用例一览", "077表")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cases = [
        ("1", "用户注册", "ENC_UC_AUTH_00", "新用户填写账号、密码、手机号完成注册"),
        ("2", "用户登录", "ENC_UC_AUTH_01", "已注册用户输入账号与密码登录, 系统下发 Sa-Token 令牌"),
        ("3", "用户登出", "ENC_UC_AUTH_02", "已登录用户主动退出, 清理令牌"),
        ("4", "浏览演出列表", "ENC_UC_BROWSE_01", "游客或注册用户浏览全部正在售票的剧目"),
        ("5", "搜索演出", "ENC_UC_BROWSE_02", "按分类、关键词、时间检索剧目"),
        ("6", "查看演出详情", "ENC_UC_BROWSE_03", "查看剧目海报、简介、时长、票价区间与场次列表"),
        ("7", "查看场次", "ENC_UC_BROWSE_04", "选择具体演出时间与演出厅, 进入选座入口"),
        ("8", "可视化选座", "ENC_UC_SEAT_01", "在座位图上选择一个或多个 AVAILABLE 座位并临时锁定"),
        ("9", "创建订单", "ENC_UC_ORDER_01", "依据已锁座位生成待支付订单并启动倒计时"),
        ("10", "取消订单", "ENC_UC_ORDER_02", "用户主动取消待支付订单, 释放座位"),
        ("11", "模拟支付", "ENC_UC_PAY_01", "用户在倒计时内完成支付模拟, 订单状态 PENDING_PAYMENT→PAID"),
        ("12", "查看我的订单", "ENC_UC_ORDER_03", "用户分页查看本人的全部订单及当前状态"),
        ("13", "查看电子票", "ENC_UC_PAY_02", "支付成功后查看票号、二维码模拟码与座位信息"),
        ("14", "演出分类管理", "ENC_UC_SHOW_00", "票务管理员维护分类树, 含新增/修改/启停"),
        ("15", "剧目信息管理", "ENC_UC_SHOW_01", "票务管理员维护剧目档案与海报"),
        ("16", "演出厅管理", "ENC_UC_HALL_01", "票务管理员维护演出厅基本信息"),
        ("17", "座位布局管理", "ENC_UC_HALL_02", "票务管理员维护演出厅座位行列与区域"),
        ("18", "场次排期管理", "ENC_UC_SCH_01", "票务管理员为剧目创建场次并选定演出厅与开演时间"),
        ("19", "场次座位票池管理", "ENC_UC_SCH_02", "基于演出厅座位生成场次专属票池, 支持票价分区"),
        ("20", "订单查询与管理", "ENC_UC_ORDER_04", "票务管理员检索和处理订单, 含异常订单"),
        ("21", "公告管理", "ENC_UC_ANN_01", "票务管理员发布、修改、下线剧院公告"),
        ("22", "查看数据看板", "ENC_UC_STAT_01", "票务管理员查看销售额、订单数、上座率、热门演出等核心指标"),
        ("23", "检票核销", "ENC_UC_CHECKIN_01", "检票员通过票号或二维码模拟码核销电子票并允许入场"),
        ("24", "异常票据处理", "ENC_UC_CHECKIN_02", "检票员对重复检票、作废票、非当前场次票据进行异常处理"),
        ("25", "用户账号管理", "ENC_UC_SYS_01", "系统管理员创建、禁用、重置用户账号"),
        ("26", "角色与权限管理", "ENC_UC_SYS_02", "系统管理员维护角色清单并为用户分配角色"),
        ("27", "操作日志查看", "ENC_UC_SYS_03", "系统管理员查询关键操作日志并导出"),
        ("28", "舞台视角预览", "ENC_UC_SEAT_02", "用户在选座页点击具体座位, 弹出 three.js 3D 场景从该座位看向舞台的视角预览"),
        ("29", "拼座邀请", "ENC_UC_SEAT_03", "用户生成可分享的拼座邀请链接 / 二维码, 邀请人扫码进入同场次联合锁座并合并下单"),
        ("30", "演出智能推荐", "ENC_UC_REC_01", "首页基于规则 + 协同过滤 + 内容相似度三路混合排序, 个性化推荐 Top 8 演出"),
    ]
    uc_tbl = insert_table_after(cur, rows=len(cases) + 1, cols=4)
    fill_row(uc_tbl.rows[0], ["序号", "用例名称", "用例标识符", "需求描述"], bold=True, center=True)
    for i, row in enumerate(cases, start=1):
        fill_row(uc_tbl.rows[i], list(row))
    cur_el = OxmlElement('w:p')
    uc_tbl._tbl.addnext(cur_el)
    cur = Paragraph(cur_el, body_wrap)
    set_style(cur, "Normal")
    cur.add_run("表 7 中共定义 30 个用例, 覆盖游客、注册用户、票务管理员、检票员、系统管理员等 5 类参与者. 其中 9 个核心用例 (ENC_UC_AUTH_01 用户登录、ENC_UC_BROWSE_01 浏览与查询演出、ENC_UC_SEAT_01 可视化选座、ENC_UC_ORDER_01 创建订单、ENC_UC_PAY_01 模拟支付、ENC_UC_CHECKIN_01 检票核销、ENC_UC_SHOW_01 剧目与场次管理、ENC_UC_STAT_01 统计报表查看、ENC_UC_REC_01 演出智能推荐) 将在 4.3 ~ 4.11 节展开详述; 实时多人选座 (ENC_UC_SEAT_01 的并发同步语义)、舞台视角预览 (ENC_UC_SEAT_02)、拼座邀请 (ENC_UC_SEAT_03) 等差异化能力作为 ENC_UC_SEAT_01 / 4.5 节的扩展场景一并描述.")

    # 表 8 数据库实体说明
    cur = insert_para_after(cur, "表8  ENCORE 数据库实体说明", "077表")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    db_entities = [
        ("sys_user", "用户账号", "保存用户登录账号、密码哈希、手机号、邮箱、状态等基本信息"),
        ("sys_role", "角色字典", "维护游客 / 注册用户 / 票务管理员 / 检票员 / 系统管理员等角色"),
        ("sys_user_role", "用户-角色关联", "多对多关联表, 实现一个用户可绑定多个角色"),
        ("show_category", "演出分类", "音乐剧 / 话剧 / 戏曲 / 综艺 等剧目分类"),
        ("show_info", "剧目信息", "剧目标题、海报、简介、时长、票价区间、所属分类"),
        ("theater_hall", "演出厅", "剧院内具体演出厅的名称、行列规模"),
        ("hall_seat", "演出厅座位", "演出厅内每个固定座位的行号、列号、区域、是否可售"),
        ("show_schedule", "演出场次", "剧目在某演出厅的具体演出时间, 携带 ScheduleStatus 状态"),
        ("schedule_seat", "场次座位票池", "依据演出厅座位为单一场次生成的可售座位副本, 携带 SeatStatus 状态与票价"),
        ("ticket_order", "订单", "用户购票订单, 携带 OrderStatus 状态、总金额、倒计时过期时间"),
        ("ticket_item", "票据明细", "支付成功后生成的电子票, 含票号、二维码模拟码、TicketStatus 状态"),
        ("payment_record", "支付记录", "记录每次模拟支付的金额、渠道、支付时间"),
        ("checkin_record", "检票记录", "记录检票员对票据的核销时间, 用于审计与防重复"),
        ("announcement", "公告", "剧院公告内容、发布时间, 用户端首页展示"),
        ("operation_log", "操作日志", "管理端关键操作日志, 含操作者、目标、动作、时间"),
    ]
    db_tbl = insert_table_after(cur, rows=len(db_entities) + 1, cols=3)
    fill_row(db_tbl.rows[0], ["数据库表名", "中文名", "用途说明"], bold=True, center=True)
    for i, row in enumerate(db_entities, start=1):
        fill_row(db_tbl.rows[i], list(row))
    cur_el = OxmlElement('w:p')
    db_tbl._tbl.addnext(cur_el)
    cur = Paragraph(cur_el, body_wrap)
    set_style(cur, "Normal")
    cur.add_run("数据库实体的关系约束: 一个演出分类下含多个剧目; 一个剧目可创建多个演出场次; 一个演出厅含多个固定座位; 一个场次基于演出厅座位生成多个 schedule_seat; 一个订单含一张或多张 ticket_item; 一张 ticket_item 唯一关联一个 schedule_seat 与至多一条 checkin_record.")

    # ============ 4.3 ~ 4.10 八大核心用例详细说明 (表 9 ~ 表 16) ============
    cur = _insert_detail_uc(cur,
        "用户注册与登录 (ENC_UC_AUTH_01)", "表9  用户注册与登录用例说明",
        _detailed_uc_rows(
            "ENC_UC_AUTH_01", "用户注册与登录", "系统用例", "用户目标级别",
            "游客 (注册时) / 注册用户 (登录时)",
            "用户: 期望快速完成账号注册并稳定登录购票; 票务管理员: 期望识别合法用户以审计追踪.",
            "用户已访问 ENCORE 用户端首页且具备可注册的手机号/邮箱.",
            "用户成功创建账号或登录, 系统在 Redis 写入 Sa-Token 会话, 返回身份信息及角色列表.",
            "1. 用户进入注册或登录页面; 2. 用户输入账号 / 密码 (注册时同步采集手机号); 3. 前端校验表单非空与格式合法; 4. 前端提交 POST /api/auth/register 或 /api/auth/login; 5. 后端校验账号唯一性 (注册) 或密码 BCrypt 校验 (登录); 6. 后端通过 Sa-Token 颁发令牌并返回身份信息; 7. 前端写入本地存储与 Pinia, 跳转主页.",
            "5a. 账号已存在 (注册场景): 提示重新选择账号. 5b. 密码错误 (登录场景): 系统返回 401, 5 次错误后短期锁定 10 分钟. 6a. 用户状态为禁用: 拒绝登录并提示联系管理员.",
            "密码须以 BCrypt 强度 ≥ 10 哈希存储; 登录接口须做防暴力破解; 登录成功 P95 响应时间 ≤ 1.0 秒.",
            "用户名: ≤ 32 字符; 密码: 6~32 字符且含字母与数字; 手机号: 11 位; Sa-Token TTL: 7 天可滚动续期; 验证码: 后续扩展, 不在本期范围.",
        ),
        doc)
    cur = insert_para_after(cur, "用户登录的关键交互见顺序图图 14.", "Normal")
    cur = insert_para_after(cur, "图14  用户注册登录顺序图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cur = _insert_detail_uc(cur,
        "浏览与查询演出 (ENC_UC_BROWSE_01)", "表10  浏览与查询演出用例说明",
        _detailed_uc_rows(
            "ENC_UC_BROWSE_01", "浏览与查询演出", "系统用例", "用户目标级别",
            "游客 / 注册用户",
            "用户: 期望快速找到感兴趣的剧目和场次; 票务管理员: 期望剧目曝光与售卖.",
            "用户已访问 ENCORE 用户端首页或演出列表页.",
            "用户成功获取剧目列表、详情或场次信息, 可继续进入选座流程.",
            "1. 用户进入演出列表页; 2. 前端调用 GET /api/show 获取剧目列表; 3. 用户按分类、关键词或时间筛选; 4. 前端调用 GET /api/show/{id} 加载详情; 5. 用户查看场次卡片; 6. 用户点击场次进入选座入口.",
            "3a. 无匹配结果: 提示 \"暂无相关演出\". 5a. 该剧目暂无 ON_SALE 场次: 仅展示场次列表但不提供选座按钮.",
            "首页与列表页须支持响应式; 首屏可见数据 P95 加载 ≤ 1.5 秒; 海报图片懒加载.",
            "分页大小: 12 条/页; 筛选条件: 分类、关键词、开演日期区间; 排序: 默认按场次开演时间升序.",
        ),
        doc)

    cur = _insert_detail_uc(cur,
        "可视化选座 (ENC_UC_SEAT_01)", "表11  可视化选座用例说明",
        _detailed_uc_rows(
            "ENC_UC_SEAT_01", "可视化选座", "系统用例", "用户目标级别",
            "注册用户",
            "用户: 期望直观看到剧院座位图并选到心仪座位; 票务管理员: 期望防重复售票, 维持高上座率; 检票员: 期望最终票据与座位严格一一对应.",
            "用户已登录, 已进入具体场次详情页, 场次状态为 ON_SALE.",
            "用户选定的座位临时锁定为 LOCKED, 系统创建状态为 PENDING_PAYMENT 的订单并启动倒计时.",
            "1. 用户进入选座页; 2. 前端 GET /api/schedule/{id}/seats 加载座位图; 3. 用户点击一个或多个 AVAILABLE 座位; 4. 前端 POST /api/schedule/{id}/lock 请求锁定; 5. 后端在 Redis SETNX 加锁并更新 schedule_seat 为 LOCKED; 6. 后端返回锁定成功; 7. 前端跳转至订单确认页.",
            "3a. 用户取消选择: 释放本地选中态. 5a. 部分座位被并发占用: 后端返回 409, 前端高亮失败座位并提示重选. 5b. 用户离线超过 30 秒未提交锁定: 前端自动放弃, 不发起锁定请求.",
            "Redis 座位锁 TTL = 15 分钟; 后端必须保证 Redis 锁与 MySQL 状态最终一致; 选座页 P95 渲染 ≤ 1.8 秒.",
            "单次最多锁定 6 个座位; 座位状态: AVAILABLE / LOCKED / SOLD / DISABLED; 锁键格式: schedule:{id}:seat:{seatId}.",
        ),
        doc)
    cur = insert_para_after(cur, "可视化选座的活动流程与系统交互分别见图 10 与图 15; 场次座位的完整状态流转见图 21.", "Normal")
    cur = insert_para_after(cur, "图10  用户选座购票活动图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur = insert_para_after(cur, "图15  可视化选座顺序图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur = insert_para_after(cur, "图21  场次座位状态图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------- 4.5 差异化扩展段 (V2 注入: WebSocket 实时 / 3D 视角 / 拼座) ----------
    cur = insert_para_after(cur, "在 ENC_UC_SEAT_01 主成功场景之上, ENCORE 票务管理系统进一步注入三项差异化扩展能力, 体现\"实时 + 沉浸 + 社交\"的剧院级选座体验:", "Normal")
    cur = insert_para_after(cur, "实时多人选座广播: 同一场次的多名在线用户通过 WebSocket / STOMP 协议订阅 /topic/schedule/{id}/seat 主题, 任一用户的锁座 / 释放动作经后端 WebSocketBroadcaster 在 P95 ≤ 200ms 内秒级推送给所有订阅者, 屏幕端无需轮询即可同步显示, 完整时序如图 25 所示.", "List Paragraph")
    cur = insert_para_after(cur, "图25  WebSocket 实时多人选座顺序图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur = insert_para_after(cur, "舞台视角预览 (ENC_UC_SEAT_02): 选座页用户点击具体座位时, 浮层基于 three.js + WebGL 渲染剧院 3D 场景, 以该座位为 PerspectiveCamera 的位置展示\"从这个座位看向舞台\"的视野预览, 帧率目标 ≥ 30 FPS, 首屏加载 ≤ 3 秒, 组件结构如图 26 所示.", "List Paragraph")
    cur = insert_para_after(cur, "图26  three.js 3D 座位舞台视角预览组件视图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur = insert_para_after(cur, "拼座邀请 (ENC_UC_SEAT_03): 用户选择 1~6 个座位并锁定后, 可生成一段 TTL = 10 分钟的拼座邀请链接 / 二维码 (含 scheduleId + 已锁 seatIds 的签名 token); 被邀请人扫码或点击链接进入同场次, 系统将其加入同一 Redis 锁集合并合并为一笔订单, 任一参与者支付即代表整组生效; 邀请超时或参与者主动退出则各自锁释放.", "List Paragraph")

    cur = _insert_detail_uc(cur,
        "创建订单 (ENC_UC_ORDER_01)", "表12  创建订单用例说明",
        _detailed_uc_rows(
            "ENC_UC_ORDER_01", "创建订单", "系统用例", "用户目标级别",
            "注册用户",
            "用户: 期望快速生成订单并在限定时间内完成支付; 票务管理员: 期望防止超卖与单据不一致.",
            "用户已成功锁定 1~6 个座位, Redis 锁键归属本用户.",
            "数据库中生成状态为 PENDING_PAYMENT 的 ticket_order 记录, 关联场次与座位, 启动订单倒计时.",
            "1. 用户在订单确认页核对座位、金额; 2. 用户点击 \"提交订单\"; 3. 前端 POST /api/order; 4. 后端校验 Redis 锁归属与有效期; 5. 后端在 MySQL 写入 ticket_order; 6. 后端在 Redis 写入 order:{id}:expire 倒计时; 7. 后端返回订单 ID; 8. 前端跳转支付页.",
            "4a. Redis 锁不存在 / 已过期 / 不属于该用户: 后端返回 412, 前端提示重新选座. 5a. 数据库写入失败: 后端回滚 Redis 锁并返回 500.",
            "订单创建接口须幂等; 倒计时窗口默认 15 分钟; 订单提交 P95 响应时间 ≤ 1.0 秒.",
            "订单 ID: 雪花算法生成; 订单金额 = Σ(scheduleSeat.price); 倒计时 TTL: 15 分钟可由后台配置.",
        ),
        doc)
    cur = insert_para_after(cur, "订单创建的关键交互见图 16; 订单完整状态流转见图 22.", "Normal")
    cur = insert_para_after(cur, "图16  创建订单顺序图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur = insert_para_after(cur, "图22  订单状态流转图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cur = _insert_detail_uc(cur,
        "模拟支付与电子票生成 (ENC_UC_PAY_01)", "表13  模拟支付与电子票生成用例说明",
        _detailed_uc_rows(
            "ENC_UC_PAY_01", "模拟支付与电子票生成", "系统用例", "用户目标级别",
            "注册用户",
            "用户: 期望在倒计时内完成支付获得电子票; 票务管理员: 期望准确记账与统计.",
            "存在状态为 PENDING_PAYMENT 且倒计时未过期的订单.",
            "订单状态 PAID, 关联座位 SOLD, 系统生成 ticket_item (UNUSED) 与 payment_record.",
            "1. 用户进入支付页; 2. 用户点击 \"模拟支付\"; 3. 前端 POST /api/payment/mock; 4. 后端校验订单状态与倒计时; 5. 后端写入 payment_record; 6. 后端更新 ticket_order 为 PAID 并将 schedule_seat 置为 SOLD; 7. 后端批量生成 ticket_item; 8. 后端返回票号列表; 9. 前端跳转 \"我的电子票\".",
            "4a. 订单已过期: 后端将订单置为 EXPIRED, 释放座位, 返回 408 \"订单已超时\". 6a. 数据库异常: 通过事务回滚, 订单与座位状态保持不变, 返回 500.",
            "支付环节为模拟实现, 不接入真实支付通道; 票号须全局唯一; 支付成功推送可用于看板实时刷新.",
            "票号 = 雪花 ID; 二维码模拟码 = ticketCode + HMAC 签名 (后续扩展); 支持的支付渠道: \"MOCK\" 单一值.",
        ),
        doc)
    cur = insert_para_after(cur, "订单支付的活动流程与系统交互分别见图 11 与图 17; 票据状态流转见图 23.", "Normal")
    cur = insert_para_after(cur, "图11  订单支付活动图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur = insert_para_after(cur, "图17  模拟支付与电子票生成顺序图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur = insert_para_after(cur, "图23  票据状态流转图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cur = _insert_detail_uc(cur,
        "检票核销 (ENC_UC_CHECKIN_01)", "表14  检票核销用例说明",
        _detailed_uc_rows(
            "ENC_UC_CHECKIN_01", "检票核销", "系统用例", "用户目标级别",
            "检票员",
            "检票员: 期望快速核验票据并放行; 票务管理员: 期望防重复检票与作废票滥用; 注册用户: 期望体验流畅.",
            "电子票存在且状态为 UNUSED; 场次时间处于允许检票窗口.",
            "票据状态 CHECKED_IN, 写入 checkin_record 记录, 允许用户入场.",
            "1. 检票员登录检票端; 2. 扫描二维码模拟码或手动输入票号; 3. 前端 POST /api/checkin/verify; 4. 后端查询 ticket_item; 5. 后端校验票据状态、场次时间; 6. 后端将票据置为 CHECKED_IN 并写入 checkin_record; 7. 前端展示核销成功与座位号.",
            "4a. 票据不存在: 后端返回 404. 5a. 票据已 CHECKED_IN: 返回 409 \"重复检票\". 5b. 票据 INVALID 或 REFUNDED: 返回 409 异常. 5c. 当前时间不在场次的允许检票窗口内: 返回 425 \"非当前场次\".",
            "异常类型均需有清晰前台提示, 防止误判; 单次核销 P95 ≤ 0.6 秒.",
            "允许检票窗口: 开场前 60 分钟至开场后 30 分钟; 检票员需要 CHECKIN 角色权限.",
        ),
        doc)
    cur = insert_para_after(cur, "检票核销的活动流程与系统交互分别见图 12 与图 18.", "Normal")
    cur = insert_para_after(cur, "图12  检票核销活动图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cur = insert_para_after(cur, "图18  检票核销顺序图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cur = _insert_detail_uc(cur,
        "剧目与场次管理 (ENC_UC_SHOW_01)", "表15  剧目与场次管理用例说明",
        _detailed_uc_rows(
            "ENC_UC_SHOW_01", "剧目与场次管理", "系统用例", "用户目标级别",
            "票务管理员",
            "票务管理员: 期望高效维护剧目和场次; 剧院运营经理: 期望排期可控; 注册用户: 期望准确的演出信息.",
            "票务管理员已登录, 拥有剧目与场次维护权限.",
            "剧目档案与场次信息已正确入库; 场次的座位票池已基于演出厅座位生成完毕.",
            "1. 管理员维护演出分类与剧目档案; 2. 管理员维护演出厅与座位布局; 3. 管理员为剧目创建场次, 选择演出厅与开演时间; 4. 系统依据演出厅座位生成 schedule_seat (全部 AVAILABLE); 5. 管理员设置场次票价分区; 6. 审核通过后将场次状态 DRAFT 改为 ON_SALE.",
            "2a. 座位布局不完整: 系统拒绝生成票池并提示完善布局. 3a. 同一演出厅时间冲突: 系统拒绝创建场次. 6a. 场次销售结束后: 自动置为 FINISHED.",
            "票池生成在事务中执行, 单次场次座位数支持至 2000; 价格分区支持 1~4 档.",
            "场次状态: DRAFT / ON_SALE / STOP_SALE / FINISHED / CANCELLED; 票价精度: 元两位小数.",
        ),
        doc)
    cur = insert_para_after(cur, "剧目与场次管理的完整活动流程见图 13.", "Normal")
    cur = insert_para_after(cur, "图13  剧目与场次管理活动图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cur = _insert_detail_uc(cur,
        "统计报表查看 (ENC_UC_STAT_01)", "表16  统计报表查看用例说明",
        _detailed_uc_rows(
            "ENC_UC_STAT_01", "统计报表查看", "系统用例", "用户目标级别",
            "票务管理员 (兼剧院运营经理只读)",
            "票务管理员: 期望实时把握销售与上座率; 剧院运营经理: 期望支持排期决策.",
            "管理员已登录后台并具有数据看板查看权限.",
            "看板中展示当日订单数、当日销售额、近 7 日票房趋势、热门演出、场次上座率等关键指标.",
            "1. 管理员打开数据看板页; 2. 前端 GET /api/stat/dashboard; 3. 后端聚合 ticket_order、ticket_item、show_schedule; 4. 后端返回 DashboardVO; 5. 前端使用 ECharts 渲染折线、柱状、饼图.",
            "2a. 当日无数据: 显示空状态而非 0 错误. 4a. 后端聚合耗时较长: 前端展示骨架屏并最长等待 3 秒.",
            "看板数据来自实时查询, 不接入离线数仓; 复杂聚合需建索引避免全表扫.",
            "时间范围: 今日 / 近 7 日 / 近 30 日; 折线点数: ≤ 30; 排行榜大小: 10 条.",
        ),
        doc)
    cur = insert_para_after(cur, "统计报表的生成交互见图 19.", "Normal")
    cur = insert_para_after(cur, "图19  统计报表生成顺序图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------- 4.10 差异化扩展段 (V2 注入: 黑金指挥中心 + 热力图 + Sankey) ----------
    cur = insert_para_after(cur, "在 ENC_UC_STAT_01 主成功场景之上, ENCORE 票务管理系统将管理端数据看板按\"黑金指挥中心\"风格实现, 不止呈现数据, 更让数据本身具有视觉冲击力. 看板架构与关键数据流如图 27 所示:", "Normal")
    cur = insert_para_after(cur, "视觉规范: 全黑底 (#0F0F0F) + 暗红 (#5B0E0E) + 金色 (#C8A35E) 高亮的深色 BI 视觉语言, 自定义 ECharts dark-encore 主题; 整体观感类似机场塔台 / 监控指挥中心, 与用户端沉浸式剧院风格遥相呼应.", "List Paragraph")
    cur = insert_para_after(cur, "高密度图表: 在传统折线 / 饼图之外, 新增 ECharts Sankey 漏斗 (浏览→选座→下单→支付, 单图表现完整转化链路)、座位区域 Heatmap (按演出厅平面坐标叠加上座率热度)、Bar Race (热门演出 Top 10 时间序列动画).", "List Paragraph")
    cur = insert_para_after(cur, "实时刷新: 后端 WebSocketBroadcaster 在用户完成支付 / 检票核销时, 通过 /topic/stat/refresh 主题推送轻量事件, 看板局部刷新关键卡片, 避免对 /api/stat/dashboard 进行高频轮询.", "List Paragraph")
    cur = insert_para_after(cur, "缓存与回源: StatAggregator 优先读取 Redis (TTL 60 秒) 命中缓存; 未命中则回源 MySQL 复杂聚合 SQL (含 GROUP BY hall_id / area / day 等), 命中率 ≥ 80% 时单次看板加载 P95 ≤ 2.5 秒.", "List Paragraph")
    cur = insert_para_after(cur, "图27  黑金\"指挥中心\"看板架构与数据流图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------- 4.11 ENC_UC_REC_01 演出智能推荐 (V2 新增子节) ----------
    cur = _insert_detail_uc(cur,
        "演出智能推荐 (ENC_UC_REC_01)", "表17  演出智能推荐用例说明",
        _detailed_uc_rows(
            "ENC_UC_REC_01", "演出智能推荐", "系统用例", "用户目标级别",
            "游客 / 注册用户",
            "用户: 期望在首页快速发现感兴趣的演出, 不必反复搜索; 票务管理员: 期望长尾剧目曝光与售卖均衡; 剧院运营经理: 期望提升整体上座率.",
            "用户访问 ENCORE 用户端首页或推荐 Tab; (注册用户场景下) 用户已登录.",
            "首页\"猜你喜欢\"位返回 Top 8 推荐演出列表, 注册用户与游客分别走个性化路径或冷启动路径, 推荐结果带相关性分数与命中策略标签.",
            "1. 用户进入首页或推荐 Tab; 2. 前端 GET /api/rec?userId=&size=8; 3. RecController 接收请求; 4. 候选生成 (Recall): 近 30 日热门 Top 50 + 用户浏览/购买分类同类热门 Top 50, 合并去重为候选池 (~80 项); 5. 三路打分 (Ranking): 协同过滤 (用户-演出共现矩阵 + 余弦相似度) + 内容过滤 (分类标签 + 时长 + 票价区间相似度) + 规则打分 (开场临近度 + 上座率 + 售罄稀缺度), 加权 0.4 / 0.3 / 0.3; 6. 排序取 Top 8; 7. 写入 recommendation:user:{id} Redis 缓存 (TTL 10 分钟); 8. 返回 RecommendVO, 前端首页\"猜你喜欢\"位渲染. 完整流程如图 28 所示.",
            "4a. 候选池为空 (新剧院冷启动): 直接走规则兜底取近 30 日热门 Top 8. 5a. 注册用户但行为不足 (用户冷启动): 关闭协同过滤路, 权重重分配为内容 0.5 / 规则 0.5. 7a. Redis 写入失败: 仍同步返回结果, 异步重试.",
            "推荐接口 P95 ≤ 1.5 秒; 协同过滤共现矩阵每日凌晨增量计算; 规则权重支持后台热配置.",
            "推荐 size: 默认 8, 上限 30; Redis 缓存 key: recommendation:user:{userId}, 匿名用户使用 IP 哈希; 后续扩展: 接入向量数据库 + 用户画像 Embedding (本期不实现).",
        ),
        doc)
    cur = insert_para_after(cur, "演出智能推荐的端到端流程见图 28.", "Normal")
    cur = insert_para_after(cur, "图28  AI 演出智能推荐流程图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    print("[Stage B-4] 第 4 章 功能需求 替换完成 (V2)")


# ============================================================
# 阶段 B-5: 第 5 章 非功能需求
# ============================================================

def stage_b_chapter5_nonfunc(doc):
    """5 非功能需求 - 完整重建 (V2 注入 WebSocket / 3D / Knife4j / 全链路压测 / DDD).
    模板 H2: 界面与接口需求 (含 H3: 界面需求 / 外部接口), 性能需求, 安全性需求,
            可靠性需求, 适应性需求, 设计约束
    """
    head = _find_h(doc, "非功能需求")
    next_h = _find_h(doc, "验收标准")
    body = doc.element.body
    body_wrap = doc._body
    el = head._p.getnext()
    while el is not None and el is not next_h._p:
        nxt = el.getnext()
        if el.tag in (qn('w:p'), qn('w:tbl')):
            body.remove(el)
        el = nxt

    cur = insert_para_after(head, "本章节描述 ENCORE 票务管理系统在功能需求之外, 必须满足的界面与接口、性能、安全、可靠性、适应性、设计约束等非功能需求, 用于约束系统在质量属性维度的设计与实现.", "Normal")

    # ---- 5.1 界面与接口需求 ----
    cur = insert_para_after(cur, "界面与接口需求", "Heading 2")
    cur = insert_para_after(cur, "界面需求", "Heading 3")
    for s in [
        "用户端 (购票端): 采用沉浸式剧院深色风格, 主色 #2A0B0B / #C8A35E / #1A1A1A, 强调海报、座位图与电子票的视觉表达; 整体布局响应式, 桌面分辨率 ≥ 1280×720, 平板与移动端通过媒体查询自适应.",
        "用户端 3D 座位视角组件: 在选座页基于 three.js 渲染剧院 3D 场景, 用户点击具体座位时弹出 \"从该座位看向舞台\" 的 PerspectiveCamera 视角预览; 帧率 ≥ 30 FPS, 首屏加载 ≤ 3 秒, WebGL 不可用时退回静态全景图.",
        "管理端: 采用 Element Plus 默认浅色商务风格, 以表格 / 表单 / 数据看板为主, 左侧固定导航 + 顶部用户菜单, 分辨率 ≥ 1366×768.",
        "管理端数据看板黑金指挥中心风格: 数据看板模块单独采用全黑底 (#0F0F0F) + 暗红 (#5B0E0E) + 金色 (#C8A35E) 高亮的深色 BI 视觉语言, 自定义 ECharts dark-encore 主题, 集成 Sankey 漏斗、座位区域热力图、Bar Race 等高密度图表, 整体观感类似机场塔台 / 监控指挥中心.",
        "检票端: 单屏极简风格, 大字体票号输入框与扫描按钮, 适配演出现场低照度环境与触控操作, 分辨率 ≥ 800×600 (横屏).",
        "整体一致性: 全部页面统一字体规范 (中文 PingFang SC / 微软雅黑, 英文 Inter), 统一栅格 (12 列, 间距 16px), 统一按钮主色与确认/危险色; 所有错误与异常提示通过 Element Plus Message 组件统一呈现.",
        "可访问性: 关键操作均提供键盘可达路径; 重要色彩对比度 ≥ WCAG AA; 所有图片与图表提供 alt / 标题描述.",
    ]:
        cur = insert_para_after(cur, s, "List Paragraph")

    cur = insert_para_after(cur, "外部接口", "Heading 3")
    for s in [
        "前后端 HTTP 接口: 全部通过 RESTful HTTP/JSON 通信, 基础路径 /api, 错误码遵循统一封装 {code, msg, data}.",
        "鉴权接口: 通过 Sa-Token 颁发与校验, 请求头携带 Authorization: Bearer <token>; 不使用 JWT (技术栈一致性).",
        "WebSocket 实时通道: 用户端 / 管理端通过 Spring WebSocket + STOMP 协议建立长连接 (/ws 端点), 后端 simp messaging 模板向 /topic/** 主题广播事件 (选座状态变更、看板局部刷新), 向 /user/queue/** 投递单播结果; 详细消息规约见表 18.",
        "接口文档: Knife4j 4.x (Swagger UI) 在 /doc.html 统一暴露全部 REST + WebSocket 接口规范, 含请求参数 / 响应模型 / 错误码示例, 支持在线调试; 生产环境通过 profile 关闭 /doc.html 路径以收敛攻击面.",
        "数据库接口: MyBatis-Plus + JDBC 连接 MySQL 8.0; 连接池采用 Spring Boot 默认 HikariCP.",
        "缓存接口: Spring Data Redis 连接 Redis 7.0, 用于 Sa-Token 会话、座位锁、订单倒计时、推荐结果缓存、看板聚合缓存.",
        "对外接口 (后续扩展): 真实支付通道、短信通道、扫码硬件本期不实现, 以接口隔离层 (xxxProvider) 预留扩展点.",
    ]:
        cur = insert_para_after(cur, s, "List Paragraph")

    # 表 18: WebSocket 消息规约
    cur = insert_para_after(cur, "ENCORE 票务管理系统的 WebSocket 消息规约 (主题路径 / 方向 / 消息体 / 触发条件) 见表 18.", "Normal")
    cur = insert_para_after(cur, "表18  ENCORE WebSocket 消息规约", "077表")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ws_rows = [
        ["主题 / 端点", "方向", "消息体示例", "触发条件 / 用途"],
        ["/ws (STOMP CONNECT)", "客户端 → 服务端", "Authorization: Bearer <token>", "建立 STOMP 连接, 携带 Sa-Token 鉴权"],
        ["/app/seat/lock", "客户端 → 服务端", "{scheduleId, seatIds:[..]}", "请求锁定一组座位"],
        ["/topic/schedule/{id}/seat", "服务端 → 全部订阅者", "{seatId, status:LOCKED|AVAILABLE|SOLD, by}", "任一用户锁座 / 释放 / 支付成功时广播, P95 ≤ 200ms"],
        ["/user/queue/lock-result", "服务端 → 单一用户", "{success, failedSeats:[..]}", "锁定请求的私有结果回执"],
        ["/topic/stat/refresh", "服务端 → 看板订阅者", "{type:ORDER_PAID|CHECKIN, payload}", "看板局部刷新触发, 避免高频轮询"],
        ["heartbeat / reconnect", "双向", "STOMP heartbeat 10s + 客户端断线指数退避重连", "保活与异常恢复; 最大重连 5 次后回退到 30s 轮询降级"],
    ]
    ws_tbl = insert_table_after(cur, rows=len(ws_rows), cols=4)
    fill_table(ws_tbl, ws_rows)
    cur_el = OxmlElement('w:p')
    ws_tbl._tbl.addnext(cur_el)
    cur = Paragraph(cur_el, body_wrap)
    set_style(cur, "Normal")
    cur.add_run("表 18 中的 6 类消息覆盖了实时多人选座 (D1) 与看板实时刷新 (D3) 两条 WebSocket 业务链路. STOMP 协议在裸 WebSocket 之上提供主题订阅与单播能力, 避免自行实现路由层.")

    # ---- 5.2 性能需求 ----
    cur = insert_para_after(cur, "性能需求", "Heading 2")
    cur = insert_para_after(cur, "ENCORE 票务管理系统的关键性能指标 (KPI) 见表 19. 指标依据中小型剧院日均订单 ≤ 1 000、并发用户 ≤ 50 的容量场景设定, 不属于真实超大型抢票场景的性能要求.", "Normal")
    cur = insert_para_after(cur, "表19  ENCORE 性能需求指标", "077表")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    perf_rows = [
        ["指标分类", "指标名称", "指标值", "测量方式 / 说明"],
        ["响应时间", "首页加载 P95", "≤ 1.5 秒", "Chrome DevTools 在 4G 模拟下首屏可见时间"],
        ["响应时间", "演出列表 P95", "≤ 1.5 秒", "GET /api/show, 12 条/页"],
        ["响应时间", "选座页渲染 P95", "≤ 1.8 秒", "GET /api/schedule/{id}/seats, 含座位图渲染"],
        ["响应时间", "登录 P95", "≤ 1.0 秒", "POST /api/auth/login"],
        ["响应时间", "下单 P95", "≤ 1.0 秒", "POST /api/order"],
        ["响应时间", "支付模拟 P95", "≤ 1.0 秒", "POST /api/payment/mock"],
        ["响应时间", "检票核销 P95", "≤ 0.6 秒", "POST /api/checkin/verify"],
        ["响应时间", "数据看板 P95", "≤ 2.5 秒", "GET /api/stat/dashboard, 含 ECharts 渲染"],
        ["响应时间", "AI 推荐 P95", "≤ 1.5 秒", "GET /api/rec, 含三路打分与排序"],
        ["实时性", "WebSocket 推送 P95", "≤ 200 ms", "选座 / 看板事件从服务端发出到全部订阅者屏幕显示"],
        ["实时性", "3D 视角加载 P95", "≤ 3 秒", "用户点击座位到 Seat3DPreview 首帧可交互"],
        ["实时性", "3D 渲染帧率", "≥ 30 FPS", "Seat3DPreview 浏览器 RAF 平均帧率"],
        ["并发", "在线并发用户", "≥ 50", "稳定承载 50 个用户同时在线浏览/选座"],
        ["并发", "下单 TPS", "≥ 10", "持续 1 分钟 10 TPS 下单, 错误率 0"],
        ["吞吐量", "检票 TPS", "≥ 20", "演出开场前 30 分钟集中检票峰值"],
        ["资源", "单实例内存", "≤ 1 GB", "Spring Boot 进程稳态 RSS"],
        ["资源", "MySQL QPS", "≤ 300", "P95 监控阈值"],
    ]
    perf_tbl = insert_table_after(cur, rows=len(perf_rows), cols=4)
    fill_table(perf_tbl, perf_rows)
    cur_el = OxmlElement('w:p')
    perf_tbl._tbl.addnext(cur_el)
    cur = Paragraph(cur_el, body_wrap)
    set_style(cur, "Normal")
    cur.add_run("表 19 中的指标通过 JMeter 5.6 全链路压测脚本回归, 覆盖 8 个核心 REST 接口 + WebSocket 长连接; Skywalking 9.x 同步抓取调用链与火焰图. 超出指标时, 优先排查 SQL 索引、N+1 查询、Redis 序列化耗时与 WebSocket 序列化策略, 必要时引入读写分离或服务多实例水平扩展.")

    # ---- 5.3 安全性需求 ----
    cur = insert_para_after(cur, "安全性需求", "Heading 2")
    for s in [
        "身份认证: 统一通过 Sa-Token 登录认证, 令牌 TTL 7 天可滚动续期, 退出登录时服务端主动失效令牌; WebSocket CONNECT 帧亦携带 Sa-Token 进行二次校验.",
        "权限控制: 基于角色 (RBAC) 与方法注解 @SaCheckRole / @SaCheckPermission 双重控制, 检票员账号不能访问票务管理或系统管理接口.",
        "密码安全: 用户密码以 BCrypt (强度 ≥ 10) 哈希存储, 数据库与日志中不出现明文密码.",
        "防暴力破解: 同一账号连续 5 次登录失败后短期锁定 10 分钟; 同一 IP 高频登录请求触发限流.",
        "数据传输: 生产环境强制 HTTPS, Nginx 配置 TLS 1.2/1.3, 禁用过时算法; WebSocket 升级为 wss://.",
        "数据校验: 前后端双重校验, 后端通过 @Validated 与全局异常处理拦截非法参数, 防止 SQL 注入、XSS、CSRF.",
        "敏感操作审计: 所有管理端关键操作 (剧目、场次、订单、用户/角色) 写入 operation_log, 含操作者、IP、时间、目标 ID.",
        "支付安全: 本期为模拟支付, 仅用于演示, 不接入真实资金通道; 后续接入时须遵循 PCI-DSS 与第三方支付平台规范.",
        "接口文档收敛: Knife4j /doc.html 仅在 dev / test profile 启用, prod profile 通过 springdoc.api-docs.enabled=false 关闭, 收敛敏感接口暴露面.",
    ]:
        cur = insert_para_after(cur, s, "List Paragraph")

    # ---- 5.4 可靠性需求 ----
    cur = insert_para_after(cur, "可靠性需求", "Heading 2")
    for s in [
        "可用性: 售票时段 (每日 09:00 ~ 23:00) 系统可用率 ≥ 99.5%, 月平均不可用时间 ≤ 3.5 小时.",
        "数据一致性: 座位锁与订单使用 Redis + MySQL 双写, 通过事务与补偿任务 (定时扫描 EXPIRED 订单) 保证最终一致.",
        "WebSocket 一致性: 客户端断线重连后须重新拉取 /api/schedule/{id}/seats 全量座位状态, 避免错过广播期间的事件; 服务端在最近 10 秒内的座位事件保留在 Redis Stream 作快速回放兜底.",
        "故障恢复: 单点应用进程崩溃可通过 systemd / Docker 自动重启; MySQL 部署主从, 故障转移 ≤ 5 分钟.",
        "备份策略: MySQL 每日 02:00 全量备份, 保留 14 天; Redis 启用 AOF 持久化, 间隔 1 秒.",
        "容错: 选座并发冲突、支付超时、检票异常、3D 渲染失败、推荐结果为空等高频边界, 均有清晰的提示与回滚 / 降级路径.",
    ]:
        cur = insert_para_after(cur, s, "List Paragraph")

    # ---- 5.5 适应性需求 ----
    cur = insert_para_after(cur, "适应性需求", "Heading 2")
    for s in [
        "浏览器适应: 用户端 / 管理端 / 检票端兼容 Chrome 100+ / Edge 100+ / Firefox 100+ / Safari 15+, 不再支持 IE.",
        "硬件适应: 用户终端可为桌面浏览器、平板浏览器、手机浏览器; 后端服务可部署于物理机、虚拟机或 Docker.",
        "3D 视角适应: Seat3DPreview 在不支持 WebGL 2.0 的设备上自动降级为静态全景图, 不阻断选座主流程.",
        "数据规模适应: 演出 ≤ 1 000、场次 ≤ 10 000、订单 ≤ 1 000 000、用户 ≤ 100 000 范围内, 列表页响应时间不退化.",
        "国际化预留: 前端文案集中管理 (Vue I18n), 未来可扩展英文; 当前版本只交付简体中文.",
        "配置适应: 端口、数据库、Redis、Sa-Token TTL、订单倒计时、座位锁 TTL、推荐权重等通过 application.yml 注入, 不必改代码.",
    ]:
        cur = insert_para_after(cur, s, "List Paragraph")

    # ---- 5.6 设计约束 (V2 改 DDD 分层) ----
    cur = insert_para_after(cur, "设计约束", "Heading 2")
    cur = insert_para_after(cur, "ENCORE 票务管理系统的设计约束在保留通用 B/S + 前后端分离 + 单一鉴权 + 统一数据库规范的基础上, 进一步引入领域驱动设计 (Domain-Driven Design, DDD) 作为后端架构方法, 让系统在功能丰富度提升后仍保持代码可演进性. 整体约束如下:", "Normal")
    for s in [
        "架构约束: 必须采用 B/S 架构与前后端分离; 后端遵循 DDD 四层分层 — interface (Controller / WebSocket Endpoint) / application (UseCase / CommandHandler) / domain (Aggregate + ValueObject + DomainService + DomainEvent) / infrastructure (MyBatis-Plus Repository / Redis / 第三方 SDK), 严格控制依赖方向 interface → application → domain ← infrastructure.",
        "限界上下文约束: 全系统划分为 7 个限界上下文 (Account / Catalog / Scheduling / Ordering / CheckIn / Analytics / Recommendation), 每个上下文拥有独立的聚合根与仓储, 通过显式领域事件 (例如 OrderPaid / SeatLocked / TicketCheckedIn) 跨上下文协作; 上下文之间禁止直接共享数据库表, 防止上下游耦合. 限界上下文清单见表 20, 拓扑结构见图 29.",
        "技术栈约束: 前端 Vue 3 + TypeScript + Element Plus + ECharts + three.js + @stomp/stompjs; 后端 Spring Boot 3 + Java 17 + MyBatis-Plus + Sa-Token + Spring WebSocket + Knife4j; 数据库 MySQL 8; 缓存 Redis 7.",
        "鉴权方案约束: 统一使用 Sa-Token, 不混用 JWT, 避免登录态实现不一致.",
        "数据库约束: 字符集 utf8mb4, 时间字段统一 DATETIME, 金额字段统一 DECIMAL(10,2), 状态字段统一 VARCHAR 枚举值, 严格遵循 limit-context-prefix 命名 (sys_ / show_ / theater_ / schedule_ / ticket_ / payment_ / checkin_).",
        "算法约束: AI 演出推荐采用规则 + 协同过滤 + 内容相似度的轻量混合方案, 不引入重型机器学习训练流程; 三路权重默认 0.4 / 0.3 / 0.3 可后台热配置, 共现矩阵每日凌晨增量计算.",
        "代码规范约束: 后端遵循阿里巴巴 Java 开发手册, 前端遵循 Airbnb Style + ESLint + Prettier 自动化校验.",
        "范围约束: 本期不实现真实支付、真实短信、真实硬件扫码、大规模抢票; AI 推荐仅做规则 + 简单协同的轻量版.",
        "授权与开源协议: 项目源代码采用 MIT 协议; 外部依赖必须为允许商用的协议 (Apache 2.0 / MIT / BSD).",
    ]:
        cur = insert_para_after(cur, s, "List Paragraph")

    # 表 20: DDD 限界上下文
    cur = insert_para_after(cur, "ENCORE 系统的 7 个限界上下文及其核心聚合与上下文映射见表 20.", "Normal")
    cur = insert_para_after(cur, "表20  ENCORE 领域驱动设计限界上下文", "077表")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ddd_rows = [
        ["限界上下文", "核心聚合 / 实体", "关键业务规则", "协作上下文 (上下文映射)"],
        ["Account 账号", "User, Role, Session", "用户身份与角色权限的唯一来源", "→ Ordering / CheckIn (Customer-Supplier 用户身份)"],
        ["Catalog 剧目", "ShowInfo, ShowCategory", "剧目元数据与分类的唯一来源", "→ Scheduling (Customer-Supplier 剧目元数据); → Recommendation (Conformist)"],
        ["Scheduling 排期", "ShowSchedule, ScheduleSeat, TheaterHall, HallSeat", "场次状态机 (DRAFT/ON_SALE/STOP_SALE/FINISHED/CANCELLED) 与座位票池生成", "→ Ordering (OHS-PL 场次座位票池); → Analytics (Conformist 上座率事件)"],
        ["Ordering 订单", "TicketOrder, TicketItem, PaymentRecord", "订单状态机 + 15 分钟倒计时 + Redis 座位锁一致性", "← Account / Scheduling (Customer-Supplier); → CheckIn / Analytics / Recommendation (Conformist 订单事件)"],
        ["CheckIn 检票", "CheckinRecord", "票据状态机 (UNUSED/CHECKED_IN/INVALID/REFUNDED) 与防重复检票", "← Account / Ordering (Customer-Supplier)"],
        ["Analytics 分析", "DashboardVO, StatAggregator", "聚合 SQL + Redis 60 秒缓存; 实时事件局部刷新看板", "← Ordering / Scheduling (Conformist)"],
        ["Recommendation 推荐", "RecommendVO, RecallEngine, RankingEngine", "三路混合排序 + 规则兜底 + 10 分钟用户级缓存", "← Catalog / Ordering (Conformist)"],
    ]
    ddd_tbl = insert_table_after(cur, rows=len(ddd_rows), cols=4)
    fill_table(ddd_tbl, ddd_rows)
    cur_el = OxmlElement('w:p')
    ddd_tbl._tbl.addnext(cur_el)
    cur = Paragraph(cur_el, body_wrap)
    set_style(cur, "Normal")
    cur.add_run("表 20 与图 29 共同描述 ENCORE 的领域分解. 上下文映射类型说明: Customer-Supplier 表示存在强业务依赖且双方协商接口; OHS-PL (Open Host Service + Published Language) 用于稳定共享的发布数据契约, 例如场次座位票池; Conformist 表示下游被动跟随上游模型, 例如分析与推荐上下文消费订单事件.")
    cur = insert_para_after(cur, "图29  ENCORE 领域驱动设计 (DDD) 限界上下文图", "077图")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    print("[Stage B-5] 第 5 章 非功能需求 替换完成 (V2)")


# ============================================================
# 阶段 B-6: 第 6 章 验收标准
# ============================================================

def stage_b_chapter6_acceptance(doc):
    """6 验收标准 - 完整重建 (V2 表 18→21, 加 AC-16~20).
    内容: 6.1 总则, 6.2 验收标准明细 (表 21 共 20 行).
    """
    head = _find_h(doc, "验收标准")
    next_h = _find_h(doc, "产品提交")
    body = doc.element.body
    body_wrap = doc._body
    el = head._p.getnext()
    while el is not None and el is not next_h._p:
        nxt = el.getnext()
        if el.tag in (qn('w:p'), qn('w:tbl')):
            body.remove(el)
        el = nxt

    cur = insert_para_after(head, "本章节定义 ENCORE 票务管理系统的验收标准, 作为项目交付时甲乙双方确认产品达标的依据.", "Normal")

    # 6.1 总则
    cur = insert_para_after(cur, "总则", "Heading 2")
    for s in [
        "验收方式: 由乙方 (ENCORE 项目组) 部署一套完整环境, 由甲方 (剧院运营方) 与课程指导教师按本章条目逐项执行验收测试.",
        "验收依据: 本《软件需求规格说明书》、配套《概要设计说明书》、《详细设计说明书》、《测试报告》、《用户手册》.",
        "通过准则: 表 21 中标记为 \"必须\" 的条目全部通过为验收通过的必要条件; \"建议\" 条目作为评分加分项, 不阻断验收.",
        "缺陷分类: 阻塞 / 严重 / 一般 / 轻微 4 个等级, 阻塞与严重缺陷必须修复并复测后才可继续验收.",
        "验收记录: 每条目须在验收报告中注明 \"通过 / 部分通过 / 不通过\" 及证据 (截图、日志、测试用例 ID).",
    ]:
        cur = insert_para_after(cur, s, "List Paragraph")

    # 6.2 验收标准明细
    cur = insert_para_after(cur, "验收标准明细", "Heading 2")
    cur = insert_para_after(cur, "ENCORE 票务管理系统的验收标准明细见表 21.", "Normal")
    cur = insert_para_after(cur, "表21  ENCORE 验收标准明细", "077表")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    acc_rows = [
        ["编号", "验收类别", "验收条目", "等级"],
        ["AC-01", "功能 - 用户", "用户可通过用户名 / 密码完成注册、登录、登出, 角色权限正确", "必须"],
        ["AC-02", "功能 - 浏览", "演出列表、搜索、详情、场次基础信息可被游客与注册用户正确访问", "必须"],
        ["AC-03", "功能 - 选座", "选座页可视化呈现座位图与四种状态, 选座并发冲突有明确提示", "必须"],
        ["AC-04", "功能 - 下单", "订单创建后 15 分钟倒计时生效, 未支付到时自动取消并释放座位", "必须"],
        ["AC-05", "功能 - 支付", "模拟支付成功后订单 PAID、座位 SOLD、电子票生成 (UNUSED), 数据一致", "必须"],
        ["AC-06", "功能 - 电子票", "用户可在 \"我的电子票\" 查看票号、二维码模拟码、座位信息", "必须"],
        ["AC-07", "功能 - 检票", "检票端可通过票号或二维码模拟码完成核销, 重复检票被正确拒绝", "必须"],
        ["AC-08", "功能 - 管理", "票务管理员可完成剧目、分类、演出厅、座位、场次、票池、订单、公告管理", "必须"],
        ["AC-09", "功能 - 看板", "数据看板可正确呈现订单数、销售额、近 7 日趋势、热门演出、上座率等指标", "必须"],
        ["AC-10", "功能 - 系统", "系统管理员可管理账号、角色、权限并查看操作日志", "必须"],
        ["AC-11", "非功能 - 性能", "表 19 中所有 P95 指标均通过 JMeter 5.6 全链路压测, 错误率 < 0.5%, Skywalking 链路截图归档", "必须"],
        ["AC-12", "非功能 - 安全", "密码 BCrypt 存储, 登录失败 5 次锁定, 操作日志完整记录, prod profile 关闭 Knife4j /doc.html", "必须"],
        ["AC-13", "非功能 - 兼容", "用户端与管理端在 Chrome / Edge / Firefox / Safari 主流版本下功能正确", "必须"],
        ["AC-14", "非功能 - 视觉", "用户端深色剧院风格一致, 关键文案与颜色无错乱, 海报与图标显示正常", "建议"],
        ["AC-15", "交付物", "源代码、文档、数据库脚本、UML 图、用户手册齐备, 项目可一键启动 (Docker Compose)", "必须"],
        ["AC-16", "差异化 - 实时", "两块屏幕分别登录不同账号, 同一场次内一方锁座, 另一方屏幕在 1 秒内同步显示 LOCKED, 不刷新页面", "必须"],
        ["AC-17", "差异化 - 3D", "选座页点击任意座位可在 3 秒内弹出 three.js 3D 舞台视角预览, 帧率 ≥ 30 FPS; WebGL 不可用时降级为静态全景图", "必须"],
        ["AC-18", "差异化 - 看板", "管理端数据看板呈现黑金指挥中心风格, 含 Sankey 漏斗、座位热力图、Bar Race, 关键卡片在订单 / 检票事件触发时实时刷新", "必须"],
        ["AC-19", "差异化 - 推荐", "首页\"猜你喜欢\"位返回 Top 8 推荐, 注册用户与游客推荐结果差异化, 冷启动时规则兜底, 接口 P95 ≤ 1.5 秒", "必须"],
        ["AC-20", "差异化 - 社交", "用户在选座页生成拼座邀请链接 / 二维码, 邀请人扫码后能加入同场次并锁同一组座位, 合并为一笔订单", "建议"],
    ]
    acc_tbl = insert_table_after(cur, rows=len(acc_rows), cols=4)
    fill_table(acc_tbl, acc_rows)
    cur_el = OxmlElement('w:p')
    acc_tbl._tbl.addnext(cur_el)
    cur = Paragraph(cur_el, body_wrap)
    set_style(cur, "Normal")
    cur.add_run("表 21 共定义 20 个验收条目 (含 V2 新增 AC-16 ~ AC-20 共 5 条差异化条目), 其中 18 条为 \"必须\" 等级, 2 条为 \"建议\" 等级. 全部 \"必须\" 条目通过即视为项目验收通过.")

    print("[Stage B-6] 第 6 章 验收标准 替换完成 (V2)")


# ============================================================
# 阶段 B-7: 第 7 章 产品提交
# ============================================================

def stage_b_chapter7_delivery(doc):
    """7 产品提交 - 完整重建 (V2 表 19→22, 加 DV-11~14).
    """
    head = _find_h(doc, "产品提交")
    next_h = _find_h(doc, "签字")
    body = doc.element.body
    body_wrap = doc._body
    el = head._p.getnext()
    while el is not None and el is not next_h._p:
        nxt = el.getnext()
        if el.tag in (qn('w:p'), qn('w:tbl')):
            body.remove(el)
        el = nxt

    cur = insert_para_after(head, "本章节列出 ENCORE 票务管理系统的最终交付物清单与提交方式.", "Normal")

    cur = insert_para_after(cur, "ENCORE 项目组应在课程设计结束前向甲方与指导教师交付如下产品, 见表 22.", "Normal")
    cur = insert_para_after(cur, "表22  ENCORE 交付物清单", "077表")
    cur.alignment = WD_ALIGN_PARAGRAPH.CENTER

    deliv_rows = [
        ["编号", "交付物名称", "形式", "说明"],
        ["DV-01", "软件需求规格说明书 (本文档)", "PDF + DOCX", f"{ENCORE_NAME_CN} SRS V1.0"],
        ["DV-02", "概要设计说明书", "PDF + DOCX", "系统模块划分、类图、接口设计"],
        ["DV-03", "详细设计说明书", "PDF + DOCX", "类与方法、SQL、关键算法详述"],
        ["DV-04", "数据库脚本", "SQL 文件", "15 张表 CREATE + 初始化数据"],
        ["DV-05", "前端源代码", "Git 仓库", "Vue 3 + TypeScript + Element Plus + ECharts + three.js"],
        ["DV-06", "后端源代码", "Git 仓库", "Spring Boot 3 + Sa-Token + MyBatis-Plus + Spring WebSocket; DDD 四层 + 7 限界上下文"],
        ["DV-07", "Docker Compose 部署脚本", "YAML 文件", "包含 Nginx + JDK + MySQL + Redis 一键启动"],
        ["DV-08", "UML 图源码与导出图", "PUML + PNG", "29 张架构 / 用例 / 活动 / 顺序 / 类 / 状态 / ER / WebSocket / 3D / 看板 / DDD 上下文图"],
        ["DV-09", "用户手册", "PDF", "面向最终用户、票务管理员、检票员、系统管理员"],
        ["DV-10", "测试用例与测试报告", "Excel + PDF", "覆盖全部 \"必须\" 验收条目"],
        ["DV-11", "Knife4j 接口文档静态导出", "HTML / PDF", "通过 Knife4j /doc.html 离线导出 REST + WebSocket 接口规范; 含请求 / 响应 / 错误码示例"],
        ["DV-12", "DDD 限界上下文与领域模型扩展文档", "PDF + DOCX", "7 个限界上下文的聚合 / 实体 / 值对象详述, 含上下文映射图与领域事件清单"],
        ["DV-13", "JMeter 全链路压测报告", "JMX + HTML + PDF", "JMeter 5.6 脚本与执行结果; 覆盖 8 个核心 REST 接口 + WebSocket 长连接; 含 P95 / 错误率 / TPS 曲线"],
        ["DV-14", "Skywalking 链路追踪截图集", "ZIP (PNG + JSON)", "Skywalking 9.x 抓取的端到端 Trace 与火焰图截图; 含选座、下单、支付、检票、看板、推荐 6 类典型场景"],
    ]
    dv_tbl = insert_table_after(cur, rows=len(deliv_rows), cols=4)
    fill_table(dv_tbl, deliv_rows)
    cur_el = OxmlElement('w:p')
    dv_tbl._tbl.addnext(cur_el)
    cur = Paragraph(cur_el, body_wrap)
    set_style(cur, "Normal")
    cur.add_run("交付方式: 全部交付物以电子件形式提交至课程指导教师指定的代码仓库与文档库, 同时刻录一份只读光盘备份归档. 文档类交付物提交 PDF + 源 DOCX 两种格式; 源代码提交 Git 仓库链接并附最终发布版本标签 (tag) v1.0.0. DV-11 ~ DV-14 为本期差异化能力 (WebSocket / 3D / DDD / AI 推荐) 配套的工程化交付物, 用于评审现场的可演示性与可验证性.")
    print("[Stage B-7] 第 7 章 产品提交 替换完成 (V2)")


# ============================================================
# 阶段 B-8: 第 8 章 签字
# ============================================================

def stage_b_chapter8_signature(doc):
    """8 签字 - 保留模板原签字表, 仅在表前补充一段说明."""
    head = _find_h(doc, "签字")
    body = doc.element.body
    body_wrap = doc._body
    # 删除 head 之后所有 w:p (保留 w:tbl)
    el = head._p.getnext()
    while el is not None:
        nxt = el.getnext()
        if el.tag == qn('w:p'):
            body.remove(el)
        el = nxt
    # 在 head 之后插入一段说明 (此段会作为签字表的引言)
    cur = insert_para_after(head, "本文档自签字之日起生效, 是 ENCORE 票务管理系统软件需求规格说明书的正式版本. 签字代表甲方 (剧院运营方) 与乙方 (ENCORE 项目组) 对本文档定义的功能需求、非功能需求、验收标准、交付物清单达成一致. 后续如需变更, 须按既定变更流程, 由双方代表共同评审并修订版本号.", "Normal")
    print("[Stage B-8] 第 8 章 签字 替换完成")


# ============================================================
# 阶段 C: 在每个 "图N 标题" 段落前插入对应 PNG
# ============================================================

def stage_c_insert_images(doc):
    """遍历 paragraphs, 每个 '图N 标题' 段落前插入一个居中图片段落."""
    import re
    inserted = 0
    missing = []
    # 用快照, 因为我们要在迭代中修改 body
    paras_snapshot = list(doc.paragraphs)
    for p in paras_snapshot:
        m = re.match(r'^图\s*(\d+)\b', p.text.strip())
        if not m:
            continue
        num = int(m.group(1))
        img_path = os.path.join(IMG_DIR, f"figure-{num:02d}.png")
        if not os.path.exists(img_path):
            missing.append(num)
            continue
        # 在 caption 之前插入图片段落
        new_p_el = OxmlElement('w:p')
        p._p.addprevious(new_p_el)
        new_p = Paragraph(new_p_el, doc._body)
        set_style(new_p, "Normal")
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_p.add_run().add_picture(img_path, width=Cm(14))
        inserted += 1
    print(f"[Stage C] 共插入 {inserted} 张图片")
    if missing:
        print(f"[Stage C] 缺失图片编号: {missing}")


# ============================================================
# 阶段 D: 一致性校验
# ============================================================

def stage_d_verify(doc):
    """校验最终 docx 的完整性与一致性, 输出报告."""
    import re
    print("\n========== [Stage D] 一致性校验 ==========")
    n_paras = len(doc.paragraphs)
    n_tables = len(doc.tables)
    n_shapes = len(doc.inline_shapes)
    print(f"段落总数: {n_paras}")
    print(f"表格总数: {n_tables}")
    print(f"图片总数 (inline_shapes): {n_shapes}")

    # 合并段落文字 + 全部表格 cell 文字
    pieces = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                pieces.append(cell.text)
    full_text = "\n".join(pieces)

    # 章节大纲
    print("\n-- 章节大纲 (Heading 1) --")
    h1s = []
    for p in doc.paragraphs:
        if p.style and p.style.name == 'Heading 1':
            h1s.append(p.text.strip())
            print(f"  {p.text.strip()}")
    expected_h1 = ["引言", "项目综述", "系统体系结构", "功能需求", "非功能需求", "验收标准", "产品提交", "签字"]
    missing_h1 = [h for h in expected_h1 if h not in h1s]
    if missing_h1:
        print(f"  ⚠ 缺失 H1: {missing_h1}")
    else:
        print("  ✅ 8 大章齐全")

    # 图编号
    print("\n-- 图编号校验 --")
    fig_nums = set()
    for m in re.finditer(r'图\s*(\d+)\b', full_text):
        fig_nums.add(int(m.group(1)))
    expected_figs = set(range(1, 30))
    missing_figs = expected_figs - fig_nums
    extra_figs = fig_nums - expected_figs
    print(f"  正文出现的图编号: {sorted(fig_nums)}")
    if missing_figs:
        print(f"  ⚠ 缺失图编号: {sorted(missing_figs)}")
    if extra_figs:
        print(f"  ⚠ 多余图编号: {sorted(extra_figs)}")
    if not missing_figs and not extra_figs:
        print("  ✅ 图 1~图 29 全部出现")

    # 表编号
    print("\n-- 表编号校验 --")
    tab_nums = set()
    for m in re.finditer(r'表\s*(\d+)\s', full_text):
        tab_nums.add(int(m.group(1)))
    expected_tabs = set(range(1, 23))
    missing_tabs = expected_tabs - tab_nums
    print(f"  正文出现的表编号: {sorted(tab_nums)}")
    if missing_tabs:
        print(f"  ⚠ 缺失表编号: {sorted(missing_tabs)}")
    else:
        print("  ✅ 表 1~表 22 全部出现")

    # 5 类参与者
    print("\n-- 参与者校验 --")
    expected_actors = ["游客", "注册用户", "票务管理员", "检票员", "系统管理员"]
    for a in expected_actors:
        cnt = full_text.count(a)
        flag = "✅" if cnt >= 2 else "⚠"
        print(f"  {flag} {a}: 出现 {cnt} 次")

    # 11 个核心用例 ID (V2 含 SEAT_02 / SEAT_03 / REC_01)
    print("\n-- 用例 ID 校验 --")
    expected_uc_ids = ["ENC_UC_AUTH_01", "ENC_UC_BROWSE_01", "ENC_UC_SEAT_01",
                       "ENC_UC_SEAT_02", "ENC_UC_SEAT_03",
                       "ENC_UC_ORDER_01", "ENC_UC_PAY_01", "ENC_UC_CHECKIN_01",
                       "ENC_UC_SHOW_01", "ENC_UC_STAT_01", "ENC_UC_REC_01"]
    for uid in expected_uc_ids:
        cnt = full_text.count(uid)
        flag = "✅" if cnt >= 1 else "⚠"
        print(f"  {flag} {uid}: {cnt} 次")

    # 4 套状态枚举
    print("\n-- 状态枚举校验 --")
    state_enums = {
        "ScheduleStatus": ["DRAFT", "ON_SALE", "STOP_SALE", "FINISHED", "CANCELLED"],
        "SeatStatus": ["AVAILABLE", "LOCKED", "SOLD", "DISABLED"],
        "OrderStatus": ["PENDING_PAYMENT", "PAID", "CANCELLED", "EXPIRED", "REFUNDED"],
        "TicketStatus": ["UNUSED", "CHECKED_IN", "INVALID", "REFUNDED"],
    }
    for name, items in state_enums.items():
        missing_items = [it for it in items if it not in full_text]
        if not missing_items:
            print(f"  ✅ {name}: 全部 {len(items)} 项出现")
        else:
            print(f"  ⚠ {name}: 缺失 {missing_items}")

    # 15 张数据库表
    print("\n-- 数据库表校验 --")
    db_tables = ["sys_user", "sys_role", "sys_user_role", "show_category", "show_info",
                 "theater_hall", "hall_seat", "show_schedule", "schedule_seat",
                 "ticket_order", "ticket_item", "payment_record", "checkin_record",
                 "announcement", "operation_log"]
    missing_db = [t for t in db_tables if t not in full_text]
    if not missing_db:
        print(f"  ✅ 15 张数据库表全部出现")
    else:
        print(f"  ⚠ 缺失: {missing_db}")

    # 鉴权一致性: Sa-Token 必须出现; 若出现 JWT 应仅为 "不使用 JWT" 这类排除说明
    print("\n-- 鉴权一致性 --")
    has_sa = "Sa-Token" in full_text
    jwt_lines = [l for l in full_text.split("\n") if "JWT" in l]
    bad_jwt = [l for l in jwt_lines if "不使用" not in l and "不混用" not in l and "Sa-Token" not in l]
    if has_sa and not bad_jwt:
        print(f"  ✅ 鉴权方案统一为 Sa-Token (JWT 仅作为排除项提及 {len(jwt_lines)} 次)")
    else:
        print(f"  ⚠ Sa-Token={has_sa}; 异常 JWT 行: {bad_jwt}")

    # V2 差异化关键词
    print("\n-- V2 差异化关键词校验 --")
    v2_keywords = ["WebSocket", "STOMP", "three.js", "DDD", "限界上下文",
                   "黑金", "Sankey", "热力图", "Knife4j", "JMeter", "Skywalking",
                   "拼座", "舞台视角", "演出智能推荐"]
    for kw in v2_keywords:
        cnt = full_text.count(kw)
        flag = "✅" if cnt >= 1 else "⚠"
        print(f"  {flag} {kw}: {cnt} 次")

    # 最低 20 张图片 (V2 期望 29)
    print("\n-- 图片插入校验 --")
    if n_shapes >= 29:
        print(f"  ✅ inline_shapes={n_shapes} (期望 29)")
    elif n_shapes >= 20:
        print(f"  △ inline_shapes={n_shapes} ≥ 20 (期望 29, 有缺失)")
    else:
        print(f"  ⚠ inline_shapes={n_shapes} < 20")
    print("==========================================\n")


# ============================================================
# Stage A 与 Stage B 入口
# ============================================================

def stage_a(doc):
    fill_cover(doc)
    fill_table_1_project(doc)
    fill_table_2_team(doc)
    fill_table_3_modify_record(doc)
    fill_table_4_terms(doc)
    fill_table_5_abbr(doc)
    fill_table_6_references(doc)
    fill_table_7_positions(doc)
    # 模板原 3x3 统计报表表 (tables[7]) 与 7x4 参与者表 (tables[8]) 不再填充,
    # 它们会在 B-2/B-4 中删除并被新表代替.
    fill_table_10_signature(doc)
    print("[Stage A] 封面 + 8 张原表 完成 (2 张待清理表跳过填充)")


def stage_b(doc):
    stage_b_chapter1_intro(doc)
    stage_b_chapter2_overview(doc)
    stage_b_chapter3_arch(doc)
    stage_b_chapter4_functional(doc)
    stage_b_chapter5_nonfunc(doc)
    stage_b_chapter6_acceptance(doc)
    stage_b_chapter7_delivery(doc)
    stage_b_chapter8_signature(doc)


if __name__ == '__main__':
    doc = Document(TEMPLATE)
    stage_a(doc)
    stage_b(doc)
    stage_c_insert_images(doc)
    stage_d_verify(doc)
    doc.save(TEMPLATE)
    print(f"保存: {TEMPLATE}")
