"""ENCORE SRS 报告填充脚本 - 第 1 部分: 工具函数与基础设置."""
from __future__ import annotations
import io
import os
import sys

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
from docx.text.paragraph import Paragraph

TEMPLATE = r"D:\ENCORE\ENCORE票务管理系统-软件需求规格说明书-V1.0.docx"
IMG_DIR = r"D:\ENCORE\uml-images"


# =============== Paragraph / Run helpers ===============

def replace_para_text(par: Paragraph, text: str):
    """替换段落文本，保留首 run 格式，删除其他 run."""
    runs = list(par.runs)
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._r.getparent().remove(r._r)
    else:
        par.add_run(text)


def clear_para(par: Paragraph):
    for r in list(par.runs):
        r._r.getparent().remove(r._r)


def set_style(par: Paragraph, style_name: str):
    try:
        doc = par.part.document
        par.style = doc.styles[style_name]
    except KeyError:
        pass


def insert_para_after(par: Paragraph, text: str = "", style_name: str | None = None,
                      align: int | None = None) -> Paragraph:
    new_p = OxmlElement('w:p')
    par._p.addnext(new_p)
    np = Paragraph(new_p, par._parent)
    if style_name:
        set_style(np, style_name)
    if text:
        np.add_run(text)
    if align is not None:
        np.alignment = align
    return np


def insert_paras_after(par: Paragraph, items: list[tuple[str, str | None]]) -> Paragraph:
    """items: [(text, style_name), ...]"""
    cur = par
    for txt, st in items:
        cur = insert_para_after(cur, txt, st)
    return cur


def remove_para(par: Paragraph):
    par._p.getparent().remove(par._p)


def insert_image_after(par: Paragraph, image_path: str, caption: str,
                       width_cm: float = 14.0) -> Paragraph:
    pic_p = insert_para_after(par, style_name='Normal')
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(image_path, width=Cm(width_cm))
    cap_p = insert_para_after(pic_p, caption, style_name='077图')
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return cap_p


def insert_table_after(par: Paragraph, rows: int, cols: int,
                       style: str = 'Table Grid') -> Table:
    doc = par.part.document
    tbl = doc.add_table(rows=rows, cols=cols)
    try:
        tbl.style = style
    except KeyError:
        pass
    tbl.autofit = True
    tbl_el = tbl._tbl
    tbl_el.getparent().remove(tbl_el)
    par._p.addnext(tbl_el)
    return tbl


# =============== Cell helpers ===============

def fill_cell(cell, text: str, bold: bool = False, center: bool = False,
              size: int | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    if bold:
        r.bold = True
    if size:
        r.font.size = Pt(size)


def fill_row(row, values: list[str], bold: bool = False, center: bool = False):
    for ci, val in enumerate(row.cells):
        if ci < len(values):
            fill_cell(val, values[ci], bold=bold, center=center)


def fill_table(table: Table, data: list[list[str]], header_bold: bool = True,
               header_center: bool = True, body_center: bool = False):
    """data[0] 是表头, 其余是数据行. 自动扩行."""
    while len(table.rows) < len(data):
        table.add_row()
    for ri, row_data in enumerate(data):
        is_header = (ri == 0)
        fill_row(
            table.rows[ri],
            row_data,
            bold=(is_header and header_bold),
            center=(is_header and header_center) or (body_center and not is_header),
        )


# =============== Paragraph location helpers ===============

def find_para_exact(doc: Document, text: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def find_para_contains(doc: Document, sub: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if sub in p.text:
            return p
    return None


def find_all_paras_contains(doc: Document, sub: str) -> list[Paragraph]:
    return [p for p in doc.paragraphs if sub in p.text]


def find_heading2_after(doc: Document, anchor: Paragraph) -> Paragraph | None:
    """从 anchor 开始向后找下一个 Heading 1 或 Heading 2."""
    seen = False
    for p in doc.paragraphs:
        if not seen:
            if p is anchor:
                seen = True
            continue
        s = p.style.name if p.style else ""
        if s in ("Heading 1", "Heading 2"):
            return p
    return None


def delete_paras_between(doc: Document, start: Paragraph, end: Paragraph | None):
    """删除 start 后到 end 前的所有 w:p 元素（保留 w:tbl）.
    start 本身不删；end 本身不删."""
    body = doc.element.body
    el = start._p.getnext()
    while el is not None and (end is None or el is not end._p):
        nxt = el.getnext()
        if el.tag == qn('w:p'):
            body.remove(el)
        el = nxt


# =============== Save export ===============

def save(doc: Document, path: str = TEMPLATE):
    doc.save(path)


if __name__ == '__main__':
    print("util module loaded")
